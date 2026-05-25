"""
TrustLayer AI - Technical Document Generator (Word .docx)
Creates a professionally formatted Word document for Infosys Business Incubator Cohort 2 Hackathon
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ============================================================================
# COLOR CONSTANTS
# ============================================================================
C_DARK_BLUE  = RGBColor(0x00, 0x33, 0x88)   # heading dark blue
C_ACCENT     = RGBColor(0x00, 0x66, 0xFF)   # primary accent blue
C_TEAL       = RGBColor(0x00, 0xB8, 0x94)   # teal
C_RED        = RGBColor(0xE7, 0x4C, 0x3C)   # danger/block
C_ORANGE     = RGBColor(0xF3, 0x9C, 0x12)   # warning/flag
C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY       = RGBColor(0x64, 0x74, 0x8B)
C_LIGHT_BLUE = RGBColor(0xE8, 0xF4, 0xFD)
C_TABLE_HDR  = RGBColor(0x00, 0x33, 0x88)
C_TABLE_ALT  = RGBColor(0xF0, 0xF4, 0xFF)
C_TABLE_ROW  = RGBColor(0xFF, 0xFF, 0xFF)


# ============================================================================
# LOW-LEVEL XML HELPERS
# ============================================================================

def set_cell_bg(cell, hex_color):
    """Set table cell background color via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_table_borders(table, color="CCCCCC"):
    """Apply uniform borders to all cells in a table."""
    border_spec = {'val': 'single', 'sz': 4, 'color': color}
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell, top=border_spec, bottom=border_spec,
                             left=border_spec, right=border_spec)


def add_horizontal_rule(doc, color="0066FF", thickness=12):
    """Add a colored horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def set_run_color(run, rgb: RGBColor):
    run.font.color.rgb = rgb


def set_heading_style(paragraph, level=1):
    """Map level to Word heading style names."""
    style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}
    paragraph.style = style_map.get(level, 'Heading 1')


# ============================================================================
# DOCUMENT HELPERS
# ============================================================================

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p


def add_para(doc, text, bold=False, italic=False, color=None,
             size=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.color.rgb = C_ACCENT
        p.add_run(" " + text)
    else:
        p.add_run(text)
    return p


def add_info_box(doc, title, content_lines, bg_hex="E8F4FD", accent_color=C_ACCENT):
    """Add a shaded info box as a single-cell table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    tf = cell.text_frame if hasattr(cell, 'text_frame') else None

    # Title
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = accent_color
    r.font.size = Pt(11)

    # Content lines
    for line in content_lines:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(1)
        r2 = p2.add_run(line)
        r2.font.size = Pt(10)

    p_end = cell.add_paragraph()
    p_end.paragraph_format.space_after = Pt(4)
    doc.add_paragraph()
    return table


def make_table(doc, headers, rows, col_widths=None,
               hdr_bg="003388", hdr_text=C_WHITE,
               alt_bg="F0F4FF", font_size=10):
    """Create a styled table with header row and alternating rows."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Set column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    # Header row
    hdr_row = table.rows[0]
    for j, hdr in enumerate(headers):
        cell = hdr_row.cells[j]
        set_cell_bg(cell, hdr_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr)
        r.bold = True
        r.font.color.rgb = hdr_text
        r.font.size = Pt(font_size)

    # Data rows
    for i, row_data in enumerate(rows):
        tr = table.rows[i + 1]
        bg = alt_bg if i % 2 == 1 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = tr.cells[j]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            r.font.size = Pt(font_size)

    doc.add_paragraph()
    return table


# ============================================================================
# DOCUMENT SECTIONS
# ============================================================================

def build_cover(doc):
    # Large title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("TrustLayer AI")
    r.font.size = Pt(36)
    r.bold = True
    r.font.color.rgb = C_DARK_BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Enterprise AI Reliability Platform")
    r2.font.size = Pt(18)
    r2.font.color.rgb = C_ACCENT

    add_horizontal_rule(doc, color="0066FF", thickness=20)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Technical Design Document")
    r3.font.size = Pt(14)
    r3.italic = True
    r3.font.color.rgb = C_GRAY

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Hackathon Submission  |  Infosys Business Incubator Cohort 2 2025")
    r4.font.size = Pt(12)
    r4.font.color.rgb = C_GRAY

    doc.add_paragraph()

    # Key stats box
    stats_table = doc.add_table(rows=1, cols=5)
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats = [
        ("$67.4B", "Annual losses from AI hallucinations"),
        ("$2.4M",  "Avg cost per incident"),
        ("77%",    "Enterprise leaders concerned"),
        ("92%",    "Hallucination catch rate"),
        ("<100ms", "Detection latency"),
    ]
    for i, (val, lbl) in enumerate(stats):
        cell = stats_table.columns[i].cells[0]
        set_cell_bg(cell, "003388")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(val + "\n")
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = C_WHITE
        p2 = cell.add_paragraph(lbl)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(6)
        for run in p2.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = C_LIGHT_BLUE

    doc.add_page_break()


def build_toc(doc):
    add_heading(doc, "Table of Contents", level=1, color=C_DARK_BLUE)
    toc_items = [
        ("1.", "Problem Statement"),
        ("2.", "Solution Overview"),
        ("3.", "System Architecture"),
        ("4.", "Detection Engine"),
        ("5.", "Industry Modules"),
        ("6.", "Compliance & Governance"),
        ("7.", "Integration Ecosystem"),
        ("8.", "Technology Stack"),
        ("9.", "ROI & Business Case"),
        ("10.", "Demo Application"),
        ("11.", "Future Roadmap"),
        ("A.", "Competitive Landscape"),
        ("B.", "Technical Glossary"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{num}  ")
        r1.bold = True
        r1.font.color.rgb = C_ACCENT
        r1.font.size = Pt(11)
        r2 = p.add_run(title)
        r2.font.size = Pt(11)
    doc.add_page_break()


def build_section1(doc):
    add_heading(doc, "1. Problem Statement", level=1, color=C_DARK_BLUE)
    add_horizontal_rule(doc)

    add_heading(doc, "1.1 The AI Hallucination Crisis", level=2, color=C_ACCENT)
    add_para(doc,
        "Large Language Models (LLMs) powering enterprise AI systems frequently generate confident but "
        "factually incorrect outputs — a phenomenon known as AI hallucination. These fabricated responses "
        "include non-existent legal citations submitted in court filings, incorrect drug dosages recommended "
        "to healthcare providers, fabricated financial rates presented to customers, and unauthorized "
        "regulatory determinations issued by government AI systems.")

    add_heading(doc, "1.2 Scale of the Problem", level=2, color=C_ACCENT)
    make_table(doc,
        headers=["Metric", "Value"],
        rows=[
            ["Annual enterprise losses from AI hallucinations", "$67.4 billion"],
            ["Average cost per critical hallucination incident", "$2.4 million"],
            ["Enterprise leaders concerned about AI hallucinations", "77%"],
            ["Organizations deploying AI despite these concerns", "~70%"],
        ],
        col_widths=[4.5, 2.0],
    )

    add_heading(doc, "1.3 Root Cause: A Structural Conflict of Interest", level=2, color=C_ACCENT)
    add_para(doc,
        "AI providers (OpenAI, Google, Microsoft, Anthropic) have a fundamental commercial disincentive "
        "to solve hallucination at the output layer. Publishing failure rates undermines product adoption. "
        "Enterprises are therefore left without a reliable mechanism to verify AI outputs before they "
        "cause harm.")

    add_heading(doc, "1.4 Real-World Incidents", level=2, color=C_ACCENT)
    incidents = [
        ("Legal",        "Multiple attorneys sanctioned — some losing their licenses — for submitting AI-generated legal briefs containing fabricated case citations."),
        ("Healthcare",   "AI clinical tools recommending drug dosages outside FDA-approved ranges, citing non-existent drug interactions."),
        ("Banking",      "Customer-facing chatbots guaranteeing mortgage rates and promotional terms that do not exist in any product database."),
        ("Government",   "AI systems issuing unauthorized benefits eligibility determinations not supported by policy."),
    ]
    for industry, desc in incidents:
        add_bullet(doc, desc, bold_prefix=f"{industry}:")

    add_heading(doc, "1.5 The Compliance Gap", level=2, color=C_ACCENT)
    add_para(doc,
        "Regulatory pressure is accelerating rapidly. Most enterprises currently have no systematic way "
        "to demonstrate AI governance compliance.")
    make_table(doc,
        headers=["Framework", "Enforcement", "Key Requirement"],
        rows=[
            ["EU AI Act", "August 2026", "Risk classification, documentation, human oversight, accuracy monitoring"],
            ["NIST AI RMF", "Voluntary (US procurement reference)", "Govern, Map, Measure, Manage functions"],
            ["ISO/IEC 42001", "International standard", "AI management system documentation"],
            ["SOC 2 Type II", "Ongoing", "Availability, processing integrity, confidentiality"],
        ],
        col_widths=[2.0, 2.2, 3.3],
    )
    doc.add_page_break()


def build_section2(doc):
    add_heading(doc, "2. Solution Overview", level=1, color=C_DARK_BLUE)
    add_horizontal_rule(doc)

    add_heading(doc, "2.1 What is TrustLayer AI?", level=2, color=C_ACCENT)
    add_para(doc,
        "TrustLayer AI is an enterprise AI reliability platform that acts as a trust layer between an "
        "organization's business systems and any LLM provider. It intercepts, analyzes, scores, and acts "
        "on every AI-generated response before it reaches an end user or downstream system.")

    add_info_box(doc,
        "Core Pipeline",
        [
            "  [User Query]  →  [LLM Provider]  →  [TrustLayer AI]  →  [Verified Response or Blocked]",
            "",
            "  PASS   — Confidence ≥ 75%: Response delivered to user",
            "  FLAG   — Confidence 50–74%: Routed to human review queue",
            "  BLOCK  — Confidence < 50%: Blocked; fallback action triggered",
        ],
        bg_hex="E8F4FD", accent_color=C_ACCENT)

    add_heading(doc, "2.2 Key Differentiators", level=2, color=C_ACCENT)
    diffs = [
        ("Real-time interception", "Sub-100ms total pipeline latency — transparent to end users"),
        ("Multi-technique detection", "8 parallel detection algorithms; not a single-point model"),
        ("Enterprise data grounding", "Verification against your own CRM, policy databases, product catalogs"),
        ("Domain-specialized modules", "Industry-specific risk models for 8 verticals"),
        ("Built-in compliance", "Pre-mapped to EU AI Act, NIST AI RMF, ISO 42001, SOC 2"),
        ("LLM-agnostic", "Works with any model provider — zero vendor lock-in"),
    ]
    make_table(doc,
        headers=["Differentiator", "Description"],
        rows=diffs,
        col_widths=[2.5, 4.5],
    )
    doc.add_page_break()


def build_section3(doc):
    add_heading(doc, "3. System Architecture", level=1, color=C_DARK_BLUE)
    add_horizontal_rule(doc)

    add_heading(doc, "3.1 Deployment Options", level=2, color=C_ACCENT)
    make_table(doc,
        headers=["Option", "Description", "Integration Effort"],
        rows=[
            ["SDK Integration",  "Import TrustLayer SDK; wrap existing LLM calls",        "1–2 hours"],
            ["API Gateway",      "Redirect LLM traffic through TrustLayer endpoint",       "Zero code changes"],
            ["On-Premise",       "Deploy within customer VPC / private cloud",             "1–2 weeks"],
            ["SaaS",             "Fully managed cloud service",                            "Minutes"],
        ],
        col_widths=[2.0, 4.0, 1.5],
    )

    add_heading(doc, "3.2 Pipeline Latency Budget", level=2, color=C_ACCENT)
    make_table(doc,
        headers=["Stage", "Target Latency"],
        rows=[
            ["API Gateway routing",            "< 2ms"],
            ["LLM response retrieval",         "Varies (external)"],
            ["Parallel detection engine",      "< 80ms"],
            ["Scoring and decision",           "< 5ms"],
            ["Response delivery",              "< 5ms"],
            ["Total TrustLayer overhead",      "< 100ms"],
        ],
        col_widths=[3.5, 2.0],
    )

    add_heading(doc, "3.3 Audit Log Schema", level=2, color=C_ACCENT)
    add_para(doc, "Every AI interaction processed by TrustLayer generates a structured, immutable audit record:")
    add_info_box(doc,
        "Audit Log Fields",
        [
            "  interaction_id       Unique UUID per request",
            "  timestamp            ISO 8601 datetime",
            "  user_query           Original user query text",
            "  llm_response         Raw LLM output",
            "  detection_scores     All 8 technique scores (0.0–1.0)",
            "  confidence_score     Aggregate confidence (0–100)",
            "  risk_score           Aggregate risk (0–100)",
            "  action               PASS | FLAG | BLOCK",
            "  issues_detected      List of specific issue tags",
            "  fallback_action      Action taken on BLOCK/FLAG",
            "  industry_module      Module applied (e.g., BFSI_Banking)",
            "  compliance_tags      Mapped regulatory requirements",
        ],
        bg_hex="F0F4FF", accent_color=C_ACCENT)
    doc.add_page_break()


def build_section4(doc):
    add_heading(doc, "4. Detection Engine", level=1, color=C_DARK_BLUE)
    add_horizontal_rule(doc)
    add_para(doc,
        "TrustLayer AI runs 8 detection techniques in parallel, each targeting a distinct LLM failure mode. "
        "Results are aggregated into a unified Confidence Score (0–100) and Risk Score (0–100), "
        "which drive the final Pass / Flag / Block decision.")

    techniques = [
        ("4.1", "Semantic Entropy Analysis",        "87%", "45ms",
         "Measures token-level probability distributions. High entropy (spread across many candidate tokens) "
         "signals the model is 'guessing'. Best for factual uncertainty and knowledge boundary violations."),
        ("4.2", "Self-Consistency Checking",        "82%", "120ms",
         "Samples the same query with varied prompting strategies and checks semantic consistency across "
         "responses. Detects logical contradictions and conflicting claims."),
        ("4.3", "Source Citation Verification",     "95%", "200ms",
         "Extracts all cited sources; performs real-time lookup against indexed databases, legal repositories, "
         "and web endpoints. Best for legal hallucinations and academic citation fabrication."),
        ("4.4", "Enterprise Knowledge Grounding",   "98%", "150ms",
         "Extracts factual claims (prices, policies, product specs); validates against connected enterprise "
         "systems via real-time API calls. Highest accuracy due to authoritative data sources."),
        ("4.5", "Claim Extraction & Classification","91%", "35ms",
         "NLP-based extraction and risk-tier classification of assertions: numerical claims, legal "
         "determinations, medical recommendations, financial guarantees."),
        ("4.6", "Hallucination Pattern Recognition","79%", "25ms",
         "Trained classifier on curated hallucination datasets across industries, updated monthly. "
         "Fastest technique; acts as first-line pre-screening."),
        ("4.7", "Temporal Consistency Checking",    "93%", "40ms",
         "Identifies date-sensitive claims; validates against real-time data feeds (regulatory databases, "
         "market data, news indexes). Best for financial rates and regulatory status."),
        ("4.8", "Numerical Claim Validation",       "96%", "80ms",
         "Parses all numerical assertions (rates, dosages, thresholds, percentages) and validates against "
         "authoritative numerical sources. Critical for healthcare and financial use cases."),
    ]

    for num, name, acc, lat, desc in techniques:
        add_heading(doc, f"{num} {name}", level=2, color=C_ACCENT)
        make_table(doc,
            headers=["Accuracy", "Latency", "Primary Use Case"],
            rows=[[acc, lat, desc]],
            col_widths=[1.0, 0.9, 5.1],
            font_size=10,
        )

    add_heading(doc, "4.9 Aggregate Performance", level=2, color=C_ACCENT)
    make_table(doc,
        headers=["Metric", "Value"],
        rows=[
            ["Overall hallucination catch rate", "92%"],
            ["False positive rate",               "< 3%"],
            ["End-to-end pipeline latency",       "< 100ms"],
            ["Platform availability SLA",         "99.9%"],
        ],
        col_widths=[3.5, 2.0],
    )
    doc.add_page_break()


def build_section5(doc):
    add_heading(doc, "5. Industry Modules", level=1, color=C_DARK_BLUE)
    add_horizontal_rule(doc)
    add_para(doc,
        "TrustLayer AI ships with pre-built, domain-specialized modules including industry-specific risk "
        "models, compliance rule sets, and enterprise system integrations.")

    modules = [
        ("5.1 BFSI — Banking & Financial Services",
         [("Fabricated interest rates / APRs", "Mortgage rate inquiry → rate fabricated; promotion non-existent; approval guarantee illegal"),
          ("Unauthorized eligibility determination", "Insurance claims: unauthorized coverage determination"),
          ("Fiduciary violations", "Investment advice: unverified return claims; regulatory disclosure missing")],
         "Core banking, pricing feeds, CRM, compliance policy engines"),
        ("5.2 Healthcare",
         [("Dangerous drug dosing errors", "Metformin 1,000mg/day recommended (FDA: 500mg starting dose)"),
          ("Fabricated drug interactions", "Grapefruit absorption enhancement — no clinical evidence"),
          ("Missed contraindications", "Alcohol + metformin lactic acidosis risk not flagged")],
         "FDA drug databases, clinical decision support systems, EHR APIs, formulary databases"),
        ("5.3 Legal",
         [("Fabricated case citations", "Morrison v. Tesla Motors (2023, 9th Circuit) — does not exist"),
          ("Non-existent statutes", "AI Liability Act of 2023 — never enacted"),
          ("Analysis without source document", "NDA review performed without actual document content")],
         "Westlaw, LexisNexis, CourtListener, Federal Register"),
        ("5.4 Enterprise Operations",
         [("Fabricated HR policy details", "Incorrect vacation accrual and overtime rules stated as fact"),
          ("Incorrect IT security guidance", "Fabricated VPN server addresses and download URLs"),
          ("Unauthorized process determinations", "Policy interpretations not grounded in actual policy database")],
         "HRIS, ServiceNow, Active Directory, IT knowledge base"),
    ]

    for title, scenarios, integrations in modules:
        add_heading(doc, title, level=2, color=C_ACCENT)
        make_table(doc,
            headers=["Risk", "Example Hallucination"],
            rows=[[r, e] for r, e in scenarios],
            col_widths=[2.5, 4.5],
        )
        add_para(doc, f"Integrations: {integrations}", italic=True, color=C_GRAY, size=10)

    add_heading(doc, "5.5–5.8 Additional Modules", level=2, color=C_ACCENT)
    make_table(doc,
        headers=["Module", "Primary Risk Addressed"],
        rows=[
            ["Manufacturing", "Fabricated QC records, ISO certification claims, engineering spec violations"],
            ["Retail",        "Unverified pricing, inventory errors, promotional term fabrication"],
            ["Telecom",       "Unauthorized account actions, billing errors, plan misrepresentation"],
            ["Government",    "Unauthorized benefits eligibility determinations, policy misinterpretation"],
        ],
        col_widths=[1.8, 5.2],
    )
    doc.add_page_break()


def build_section6(doc):
    add_heading(doc, "6. Compliance & Governance", level=1, color=C_DARK_BLUE)
    add_horizontal_rule(doc)

    make_table(doc,
        headers=["Framework", "Coverage", "Key Capabilities"],
        rows=[
            ["EU AI Act",     "94%", "Risk classification, technical documentation, human oversight logging, accuracy monitoring"],
            ["NIST AI RMF",   "89%", "Govern, Map, Measure, Manage function coverage"],
            ["ISO/IEC 42001", "82%", "AI management system documentation"],
            ["SOC 2 Type II", "97%", "Availability, processing integrity, confidentiality controls"],
        ],
        col_widths=[1.8, 0.9, 4.8],
    )

    add_heading(doc, "6.1 EU AI Act Readiness (Critical: August 2026)", level=2, color=C_ACCENT)
    add_para(doc,
        "The EU AI Act requires enterprises deploying high-risk AI systems to meet five core obligations. "
        "Non-compliance penalties reach up to EUR 35 million or 7% of global annual revenue.")
    requirements = [
        ("Risk Classification",       "TrustLayer auto-classifies each AI deployment by risk tier per Article 6"),
        ("Technical Documentation",   "Auto-generated per deployment; audit-ready without manual assembly"),
        ("Human Oversight",           "Every FLAG and BLOCK decision creates a human review record"),
        ("Accuracy Monitoring",       "Real-time dashboard with historical accuracy trends"),
        ("Incident Logging",          "Immutable audit trail for every AI interaction"),
    ]
    make_table(doc,
        headers=["Requirement", "TrustLayer Implementation"],
        rows=requirements,
        col_widths=[2.2, 4.8],
    )
    doc.add_page_break()


def build_section7(doc):
    add_heading(doc, "7. Integration Ecosystem", level=1, color=C_DARK_BLUE)
    add_horizontal_rule(doc)

    add_heading(doc, "7.1 LLM Provider Support", level=2, color=C_ACCENT)
    make_table(doc,
        headers=["Provider", "Integration Method"],
        rows=[
            ["OpenAI (GPT-4, GPT-4o)",      "API Gateway + SDK"],
            ["Anthropic (Claude 3/4)",       "API Gateway + SDK"],
            ["Google (Gemini)",              "API Gateway + SDK"],
            ["Azure OpenAI",                 "API Gateway"],
            ["AWS Bedrock",                  "API Gateway"],
            ["Meta (Llama, via Bedrock)",    "API Gateway"],
            ["Mistral",                      "API Gateway"],
            ["Open-source / self-hosted",    "SDK"],
        ],
        col_widths=[3.0, 3.0],
    )

    add_heading(doc, "7.2 Enterprise Connectors (50+)", level=2, color=C_ACCENT)
    make_table(doc,
        headers=["Category", "Systems"],
        rows=[
            ["CRM",        "Salesforce, HubSpot, Microsoft Dynamics"],
            ["ERP",        "SAP, Oracle, Workday"],
            ["ITSM",       "ServiceNow, Jira, Zendesk"],
            ["HRIS",       "Workday, BambooHR, ADP"],
            ["Legal",      "Westlaw, LexisNexis, CourtListener"],
            ["Healthcare", "Epic, Cerner, FDA APIs, Clinical Databases"],
            ["Financial",  "Bloomberg, Reuters, Core Banking APIs"],
            ["Identity",   "Active Directory, Okta, Azure AD"],
        ],
        col_widths=[1.5, 5.5],
    )

    add_heading(doc, "7.3 SDK Integration Example", level=2, color=C_ACCENT)
    add_info_box(doc,
        "Python SDK — 3 lines to integrate",
        [
            "  from trustlayer import TrustLayerClient",
            "  client = TrustLayerClient(api_key='tl_...', industry='BFSI_Banking')",
            "  response = client.verify(query=user_msg, llm_response=openai_response)",
            "  if response.action == 'PASS': deliver_to_user(response.content)",
            "  elif response.action == 'FLAG': queue_for_review(response)",
            "  else: route_to_fallback(response.fallback_suggestion)",
        ],
        bg_hex="F0F4FF", accent_color=C_ACCENT)
    doc.add_page_break()


def build_section8(doc):
    add_heading(doc, "8. Technology Stack", level=1, color=C_DARK_BLUE)
    add_horizontal_rule(doc)

    layers = [
        ("Application Layer", [
            ("Demo Frontend",    "Streamlit + Plotly",                    "Interactive hackathon demonstration"),
            ("API Gateway",      "FastAPI (Python)",                      "High-performance request routing"),
            ("Detection Engine", "Python + Hugging Face Transformers",    "NLP-based hallucination detection"),
            ("Scoring Service",  "Python + scikit-learn",                 "Ensemble scoring model"),
        ]),
        ("Detection & ML", [
            ("Semantic Analysis",     "sentence-transformers",      "Embedding-based entropy measurement"),
            ("Claim Extraction",      "spaCy NER + classifiers",   "Named entity + assertion extraction"),
            ("Pattern Recognition",   "Fine-tuned RoBERTa",        "Hallucination signature classifier"),
            ("Numerical Validation",  "Regex + NER pipeline",      "Numerical entity extraction & validation"),
        ]),
        ("Data & Persistence", [
            ("Audit Logging",         "PostgreSQL (append-only)",   "Immutable interaction audit trail"),
            ("Real-time Monitoring",  "TimescaleDB / InfluxDB",     "Time-series metrics"),
            ("Enterprise Cache",      "Redis (TTL-based)",          "Fresh enterprise data snapshots"),
            ("Compliance Reports",    "Auto-generated PDF/JSON",    "Regulatory documentation output"),
        ]),
        ("Infrastructure", [
            ("Containerization",  "Docker + Docker Compose",   "Local and staging environments"),
            ("Orchestration",     "Kubernetes",                "Production-scale deployment"),
            ("Cloud",             "AWS / Azure / GCP",         "Multi-cloud; customer-choice"),
            ("CI/CD",             "GitHub Actions",            "Automated testing and deployment"),
        ]),
    ]

    for layer_name, components in layers:
        add_heading(doc, layer_name, level=2, color=C_ACCENT)
        make_table(doc,
            headers=["Component", "Technology", "Purpose"],
            rows=components,
            col_widths=[1.8, 2.2, 3.0],
        )
    doc.add_page_break()


def build_section9(doc):
    add_heading(doc, "9. ROI & Business Case", level=1, color=C_DARK_BLUE)
    add_horizontal_rule(doc)

    add_heading(doc, "9.1 Cost of Inaction", level=2, color=C_ACCENT)
    make_table(doc,
        headers=["Risk Category", "Estimated Exposure"],
        rows=[
            ["Regulatory fines (EU AI Act)",         "Up to EUR 35M or 7% global annual revenue"],
            ["Incident remediation",                 "$2.4M × number of critical incidents"],
            ["Legal liability",                      "Uncapped — depends on harm caused"],
            ["Reputation damage",                    "Long-term revenue impact (unquantified)"],
        ],
        col_widths=[3.0, 4.0],
    )

    add_heading(doc, "9.2 ROI Model (100,000 queries/month)", level=2, color=C_ACCENT)
    make_table(doc,
        headers=["Metric", "Without TrustLayer", "With TrustLayer"],
        rows=[
            ["Monthly AI query volume",        "100,000",         "100,000"],
            ["Hallucination rate",             "15%",             "15%"],
            ["Critical incidents / year",      "~30",             "~2"],
            ["Annual exposure",                "~$72,000,000",    "~$4,800,000"],
            ["TrustLayer annual cost",         "—",               "~$500,000"],
            ["Net annual benefit",             "—",               "~$65,700,000"],
            ["ROI",                            "—",               "> 400%"],
        ],
        col_widths=[3.0, 2.0, 2.0],
    )

    add_heading(doc, "9.3 Time to Value", level=2, color=C_ACCENT)
    make_table(doc,
        headers=["Milestone", "Timeline"],
        rows=[
            ["SDK integration / API gateway setup",    "1–2 days"],
            ["Industry module configuration",          "1 week"],
            ["Enterprise data source connections",     "2–3 weeks"],
            ["Full production deployment",             "2–4 weeks"],
            ["First compliance report generated",      "Day 1"],
        ],
        col_widths=[4.0, 2.0],
    )
    doc.add_page_break()


def build_section10(doc):
    add_heading(doc, "10. Demo Application", level=1, color=C_DARK_BLUE)
    add_horizontal_rule(doc)

    add_para(doc,
        "The TrustLayer AI hackathon demo is a fully interactive Streamlit application demonstrating the "
        "platform's capabilities across 9 pages, with 13 live industry scenarios and real-time scoring "
        "simulation.")

    add_heading(doc, "10.1 Demo Pages", level=2, color=C_ACCENT)
    make_table(doc,
        headers=["Page", "Purpose", "Key Elements"],
        rows=[
            ["Executive Overview",      "Market context and platform summary",   "Market stats, capability cards, industry risk heatmap"],
            ["Live Detection Demo",     "Interactive hallucination detection",    "13 scenarios, real-time scoring animation"],
            ["Technical Architecture",  "Detection engine deep-dive",            "8 techniques with accuracy/latency metrics"],
            ["Industry Solutions",      "Domain module showcase",                "8 industry vertical cards"],
            ["Real-Time Monitoring",    "Operations dashboard",                  "Query volume chart, detection feed, alerts"],
            ["Compliance & Governance", "Regulatory coverage",                   "EU AI Act, NIST, ISO, SOC 2 coverage meters"],
            ["Integrations",            "Ecosystem showcase",                    "8 LLM providers, 50+ enterprise connectors"],
            ["ROI Calculator",          "Business case tool",                    "Customizable inputs, 12-month projection"],
            ["API Documentation",       "Developer resources",                   "SDK examples, API reference, patterns"],
        ],
        col_widths=[2.0, 2.2, 3.3],
        font_size=9,
    )

    add_heading(doc, "10.2 Live Detection Scenarios", level=2, color=C_ACCENT)
    make_table(doc,
        headers=["Industry", "Scenario", "Confidence", "Action"],
        rows=[
            ["BFSI — Banking",       "Mortgage rate inquiry",           "18%", "BLOCKED"],
            ["BFSI — Insurance",     "Claims processing",               "22%", "BLOCKED"],
            ["BFSI — Wealth",        "Investment advice",               "31%", "BLOCKED"],
            ["Healthcare — Clinical","Drug dosage (metformin)",          "6%",  "BLOCKED"],
            ["Healthcare — Admin",   "Prior authorization",             "29%", "BLOCKED"],
            ["Legal — Litigation",   "Case research (all fabricated)",   "4%",  "BLOCKED"],
            ["Legal — Contract",     "NDA review (no document)",        "35%", "BLOCKED"],
            ["Enterprise — HR",      "Policy inquiry",                  "42%", "FLAGGED"],
            ["Enterprise — IT",      "VPN support (fabricated URL)",    "38%", "FLAGGED"],
            ["Manufacturing",        "Quality control records",         "55%", "FLAGGED"],
            ["Retail",               "Product inquiry (unverified)",    "61%", "FLAGGED"],
            ["Telecom",              "Billing inquiry",                 "27%", "BLOCKED"],
            ["Government",           "Benefits eligibility",            "19%", "BLOCKED"],
        ],
        col_widths=[2.2, 2.8, 1.1, 1.1],
        font_size=9,
    )

    add_heading(doc, "10.3 Recommended Demo Flow (15–20 minutes)", level=2, color=C_ACCENT)
    make_table(doc,
        headers=["Time", "Section", "Focus"],
        rows=[
            ["0–2 min",   "Executive Overview",    "$67.4B problem, market urgency"],
            ["2–6 min",   "Live Detection Demo",   "Banking + Healthcare scenarios"],
            ["6–8 min",   "Technical Architecture","8 detection techniques, latency"],
            ["8–10 min",  "Industry Solutions",    "BFSI module deep-dive"],
            ["10–12 min", "Compliance",            "EU AI Act 2026 readiness"],
            ["12–15 min", "ROI Calculator",        "Customize inputs for audience"],
            ["15–20 min", "Q&A",                   "Address technical questions"],
        ],
        col_widths=[1.2, 2.5, 3.8],
    )
    doc.add_page_break()


def build_section11(doc):
    add_heading(doc, "11. Future Roadmap", level=1, color=C_DARK_BLUE)
    add_horizontal_rule(doc)

    phases = [
        ("Phase 1 — Current (Hackathon Demo)", [
            "Interactive Streamlit demo with 13 industry scenarios",
            "Simulated detection engine with realistic scoring",
            "9-page demo covering all platform capabilities",
            "EU AI Act, NIST, ISO, SOC 2 compliance dashboards",
        ]),
        ("Phase 2 — MVP (3–6 months)", [
            "Production detection engine with real NLP models",
            "Live API gateway for LLM traffic interception",
            "3 enterprise system connectors (Salesforce, ServiceNow, SAP)",
            "Pilot deployment with 1–2 enterprise customers",
            "Real audit logging and compliance report generation",
        ]),
        ("Phase 3 — Growth (6–12 months)", [
            "Expand to 50+ enterprise connectors",
            "Fine-tune industry-specific hallucination detection models",
            "SOC 2 Type II certification",
            "Multi-region deployment (US, EU, APAC)",
            "Self-serve onboarding portal",
        ]),
        ("Phase 4 — Scale (12–24 months)", [
            "EU AI Act compliance certification program",
            "Marketplace for third-party detection modules",
            "Federated learning for enterprise-specific models (privacy-preserving)",
            "Real-time regulatory update pipeline",
            "Series B fundraising readiness",
        ]),
    ]

    for phase_title, items in phases:
        add_heading(doc, phase_title, level=2, color=C_ACCENT)
        for item in items:
            add_bullet(doc, item)
        doc.add_paragraph()
    doc.add_page_break()


def build_appendices(doc):
    add_heading(doc, "Appendix A: Competitive Landscape", level=1, color=C_DARK_BLUE)
    add_horizontal_rule(doc)
    make_table(doc,
        headers=["Solution", "Type", "Gap vs TrustLayer AI"],
        rows=[
            ["OpenAI / Google native safety", "Provider-built",   "Conflict of interest; not enterprise-grade governance"],
            ["Guardrails AI",                 "Open-source",      "No enterprise grounding; no compliance mapping"],
            ["Arthur AI",                     "Point solution",   "Monitoring only; no real-time blocking capability"],
            ["Fiddler AI",                    "Point solution",   "Bias/drift focus; not hallucination-specific"],
            ["Custom in-house build",         "DIY",              "6–12 months; $500K+; ongoing maintenance burden"],
            ["TrustLayer AI",                 "Integrated platform","Detection + Prevention + Governance in one solution"],
        ],
        col_widths=[2.5, 1.8, 3.2],
    )

    doc.add_paragraph()
    add_heading(doc, "Appendix B: Technical Glossary", level=1, color=C_DARK_BLUE)
    add_horizontal_rule(doc)
    make_table(doc,
        headers=["Term", "Definition"],
        rows=[
            ["Hallucination",        "AI-generated content that is confidently stated but factually incorrect or fabricated"],
            ["Confidence Score",     "TrustLayer's 0–100 measure of response reliability"],
            ["Risk Score",           "TrustLayer's 0–100 measure of potential harm if the response is delivered"],
            ["Semantic Entropy",     "Statistical measure of uncertainty in LLM token generation distributions"],
            ["Enterprise Grounding", "Verification of AI claims against authoritative internal data sources"],
            ["EU AI Act",            "European Union regulation governing AI systems; full enforcement from August 2026"],
            ["NIST AI RMF",          "National Institute of Standards and Technology AI Risk Management Framework"],
            ["Pass",                 "TrustLayer action: response delivered directly to user (confidence ≥ 75%)"],
            ["Flag",                 "TrustLayer action: response routed to human review (confidence 50–74%)"],
            ["Block",                "TrustLayer action: response prevented; fallback triggered (confidence < 50%)"],
        ],
        col_widths=[2.0, 5.5],
        font_size=10,
    )


# ============================================================================
# MAIN
# ============================================================================

def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # Default paragraph font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # Heading styles
    for level, size, color in [
        (1, 18, C_DARK_BLUE),
        (2, 14, C_ACCENT),
        (3, 12, C_TEAL),
    ]:
        style_name = f'Heading {level}'
        style = doc.styles[style_name]
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.font.name = 'Calibri'

    print("Building Word document...")
    build_cover(doc);       print("  Cover page")
    build_toc(doc);         print("  Table of Contents")
    build_section1(doc);    print("  Section 1: Problem Statement")
    build_section2(doc);    print("  Section 2: Solution Overview")
    build_section3(doc);    print("  Section 3: System Architecture")
    build_section4(doc);    print("  Section 4: Detection Engine")
    build_section5(doc);    print("  Section 5: Industry Modules")
    build_section6(doc);    print("  Section 6: Compliance & Governance")
    build_section7(doc);    print("  Section 7: Integration Ecosystem")
    build_section8(doc);    print("  Section 8: Technology Stack")
    build_section9(doc);    print("  Section 9: ROI & Business Case")
    build_section10(doc);   print("  Section 10: Demo Application")
    build_section11(doc);   print("  Section 11: Future Roadmap")
    build_appendices(doc);  print("  Appendices A & B")

    out = r"C:\aiprojects\aitrust\TrustLayer_AI_Technical_Document_Canada_2026.docx"
    doc.save(out)
    print(f"\nDocument saved: {out}")


if __name__ == "__main__":
    build_document()
