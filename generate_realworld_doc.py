from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def rgb(h):
    return RGBColor(int(h[:2], 16), int(h[2:4], 16), int(h[4:], 16))


def heading(doc, text, level=1, color="1F3864"):
    h = doc.add_heading(text, level=level)
    if h.runs:
        h.runs[0].font.color.rgb = rgb(color)
    return h


def para(doc, text, size=11, bold=False, italic=False, color=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = rgb(color)
    return p


def bullet(doc, items, size=11, indent=0.3):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(indent)
        r = p.add_run(item)
        r.font.size = Pt(size)


def numbered(doc, items, size=11):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        r.font.size = Pt(size)


def table(doc, headers, rows, hdr_color, col_widths=None):
    row_colors = ["F0F4F8", "FFFFFF"]
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = h
        set_cell_bg(c, hdr_color)
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = rgb("FFFFFF")
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        rc = tbl.add_row().cells
        bg = row_colors[ri % 2]
        for ci, val in enumerate(row):
            rc[ci].text = val
            set_cell_bg(rc[ci], bg)
            for p in rc[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[ci].width = Inches(w)
    return tbl


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.9)
        s.bottom_margin = Inches(0.9)
        s.left_margin = Inches(1.1)
        s.right_margin = Inches(1.1)

    # ── Cover ──────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("TrustLayer AI")
    r.bold = True
    r.font.size = Pt(32)
    r.font.color.rgb = rgb("1F3864")

    s2 = doc.add_paragraph()
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s2.add_run("Real-World Deployment & Plain-English Explanation Guide")
    r2.font.size = Pt(16)
    r2.font.color.rgb = rgb("2E75B6")

    doc.add_paragraph()
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.add_run(
        "How TrustLayer fits into your existing AI stack\n"
        "and how to explain it simply to any audience\n"
        "Canada Hackathon 2026  |  Enterprise AI Governance"
    ).font.size = Pt(11)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 1 — THE ONE-LINE EXPLANATION
    # ══════════════════════════════════════════════════════════
    heading(doc, "1.  The One-Line Explanation", color="1F3864")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        '"TrustLayer is a fact-checker that sits between your AI and your users '
        '— it stops wrong answers before they reach anyone."'
    )
    r.bold = True
    r.italic = True
    r.font.size = Pt(14)
    r.font.color.rgb = rgb("2E75B6")

    doc.add_paragraph()
    para(doc,
        "That is the sentence to use with any non-technical audience. Everything else in this "
        "document is an expansion of that one idea — for different stakeholders, at different levels "
        "of depth."
    )

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 2 — THE REAL PROBLEM IN PLAIN ENGLISH
    # ══════════════════════════════════════════════════════════
    heading(doc, "2.  The Real Problem in Plain English", color="1F3864")
    para(doc,
        "When you deploy an AI chatbot — in a bank, a hospital, a law firm — "
        "the AI sometimes makes things up. Not randomly. It makes things up confidently, "
        "fluently, and in a way that looks completely correct."
    )
    doc.add_paragraph()
    para(doc, "Here is what that looks like in banking:", bold=True)
    bullet(doc, [
        "Customer asks: 'What is your current mortgage rate?'",
        "AI answers: '4.2% fixed for 30 years — and you are pre-approved!'",
        "Reality: The actual rate is 6.85%. There is no pre-approval. No promotion exists.",
        "Customer calls a branch agent expecting 4.2%. Bank faces a regulatory complaint.",
    ])
    doc.add_paragraph()
    para(doc,
        "The AI did not lie maliciously. It generated the most statistically plausible-sounding "
        "answer. This is called a hallucination. And there is currently no standard enterprise "
        "tool that catches it in real time, before the customer reads it."
    )
    para(doc, "TrustLayer is that tool.", bold=True)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 3 — WHERE IT SITS IN THE REAL WORLD
    # ══════════════════════════════════════════════════════════
    heading(doc, "3.  Where TrustLayer Sits in a Real Enterprise Stack", color="1F3864")

    para(doc,
        "TrustLayer is a middleware intercept layer. It does not replace your AI. "
        "It does not replace your application. It plugs in between them at one point "
        "and acts as a gate on every response before the user sees it."
    )

    doc.add_paragraph()
    heading(doc, "3.1  Without TrustLayer (today)", level=2, color="CC2222")
    bullet(doc, [
        "Customer submits a question through your banking app or website chatbot.",
        "Your app sends the question to an AI model (Claude, GPT-4o, Gemini, etc.).",
        "The AI generates a response.",
        "The response goes directly to the customer. No verification. No check.",
        "If the response contains a fabricated rate, a false guarantee, or wrong eligibility — the customer reads it.",
    ])

    doc.add_paragraph()
    heading(doc, "3.2  With TrustLayer (the new flow)", level=2, color="007A3D")
    bullet(doc, [
        "Customer submits a question through your banking app or website chatbot.",
        "Your app sends the question to your AI model.",
        "The AI generates a response.",
        "TrustLayer intercepts that response BEFORE it reaches the customer.",
        "TrustLayer runs 8 detection checks against authoritative grounding data in under 2 seconds.",
        "PASS: response is clean — delivered to customer immediately.",
        "FLAG: response is uncertain — routed to a human compliance reviewer.",
        "BLOCK: response contains a dangerous or false claim — customer receives a safe fallback message instead.",
    ])

    doc.add_paragraph()
    heading(doc, "3.3  The Three Deployment Patterns", level=2, color="2E75B6")
    para(doc,
        "TrustLayer can be integrated into any enterprise AI stack in three ways:"
    )
    doc.add_paragraph()

    table(doc,
        ["Pattern", "How it works", "Best for", "Integration effort"],
        [
            ("API Gateway",
             "TrustLayer is called as a REST API between your app and the LLM. "
             "Your app sends: {query, response, industry}. TrustLayer returns: {decision, score, issues}.",
             "Any existing chatbot or AI product regardless of tech stack",
             "Low — one API call added to existing code"),
            ("SDK / Library",
             "TrustLayer Python package imported directly into your application. "
             "detector.analyze() called in-process before returning the response to the user.",
             "Python-based backends, Streamlit apps, FastAPI services",
             "Very low — 3 lines of code"),
            ("Reverse Proxy / Sidecar",
             "TrustLayer deployed as a sidecar container or nginx module. "
             "All traffic to your LLM endpoint is automatically intercepted and scored. "
             "No application code change required.",
             "Containerised enterprise deployments, Kubernetes, AWS ECS",
             "Medium — infrastructure config only, zero app changes"),
        ],
        "1F3864",
        col_widths=[1.3, 2.6, 1.8, 1.5]
    )

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 4 — EXPLAINING INFERENCE SIMPLY
    # ══════════════════════════════════════════════════════════
    heading(doc, "4.  Explaining 'Inference' Without the Jargon", color="1F3864")
    para(doc,
        "The word 'inference' confuses most non-technical stakeholders. Do not use it. "
        "Below are audience-specific translations of what TrustLayer does — "
        "each is accurate and uses zero technical jargon."
    )
    doc.add_paragraph()

    table(doc,
        ["Audience", "Plain-English Explanation"],
        [
            ("Banking customer",
             "Every answer you receive from our AI assistant is fact-checked against our "
             "real product data before it reaches you. If the AI makes something up, "
             "you never see it."),
            ("Branch manager / frontline staff",
             "Think of TrustLayer as a compliance officer sitting between the chatbot and "
             "the customer. It reads every answer, compares it to our actual rates and "
             "policies, and stops wrong answers before they cause a problem."),
            ("Compliance / risk officer",
             "TrustLayer is a real-time automated pre-delivery review layer. Every AI "
             "response is scored against authoritative product data and regulatory rules "
             "before reaching the customer. Flagged responses are routed to human review "
             "with a full audit trail."),
            ("CTO / IT leadership",
             "TrustLayer is a middleware API that intercepts LLM responses, runs structured "
             "multi-technique scoring against injected grounding data, and returns a "
             "PASS / FLAG / BLOCK decision with a confidence score before the response "
             "is delivered. It is model-agnostic and adds 1 to 3 seconds latency."),
            ("Regulator / auditor",
             "TrustLayer provides automated, logged, policy-driven oversight of every AI "
             "output before customer exposure. It produces an auditable record of each "
             "decision, the reasoning behind it, and the specific claims that were flagged "
             "or blocked. It maps to EU AI Act Article 9 risk management requirements."),
            ("CEO / board",
             "We added a safety layer to our AI assistant that catches wrong answers "
             "before customers see them. If the AI invents a rate or makes a promise "
             "we cannot keep, TrustLayer blocks it automatically. "
             "It is the seatbelt on our AI product."),
        ],
        "1F3864",
        col_widths=[1.8, 5.0]
    )

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 5 — INDUSTRY-SPECIFIC REAL WORLD USE CASES
    # ══════════════════════════════════════════════════════════
    heading(doc, "5.  Real-World Use Cases by Industry", color="1F3864")
    para(doc,
        "TrustLayer is not limited to banking. Here is how it maps to each vertical "
        "and the specific harm it prevents in each."
    )
    doc.add_paragraph()

    heading(doc, "5.1  Banking & Financial Services (BFSI)", level=2, color="2E75B6")
    table(doc,
        ["Scenario", "Without TrustLayer", "With TrustLayer"],
        [
            ("Customer asks for current mortgage rate",
             "AI returns 4.2%. Actual rate is 6.85%. Customer visits branch expecting 4.2%. Complaint filed.",
             "Enterprise Grounding detects rate deviation. Response BLOCKED. Customer receives: 'Please speak to an advisor for current rates.'"),
            ("Customer asks if they are pre-approved",
             "AI says 'Yes, you are pre-approved for $400,000.' Legally misleading. Regulatory breach.",
             "Pattern Recognition detects illegal guarantee language. Response FLAGGED for compliance review."),
            ("Customer asks about a promotion",
             "AI invents a cashback promotion. Customer expects the promotion at application. Bank has no record of it.",
             "Grounding confirms no promotions active. Fabricated promotion BLOCKED before customer reads it."),
        ],
        "2E75B6",
        col_widths=[1.8, 2.4, 2.6]
    )

    doc.add_paragraph()
    heading(doc, "5.2  Healthcare", level=2, color="CC2222")
    table(doc,
        ["Scenario", "Without TrustLayer", "With TrustLayer"],
        [
            ("Patient asks about Metformin starting dose",
             "AI recommends 1000mg. FDA grounding says correct starting dose is 500mg. Patient safety risk.",
             "Numerical Validation detects wrong dosage against FDA grounding. Response BLOCKED."),
            ("Patient asks if ibuprofen and warfarin can be combined",
             "AI says yes with a recommended dose. These drugs are CONTRAINDICATED. Serious bleeding risk.",
             "Enterprise Grounding flags CONTRAINDICATED drug combination. Response BLOCKED immediately."),
        ],
        "CC2222",
        col_widths=[1.8, 2.4, 2.6]
    )

    doc.add_paragraph()
    heading(doc, "5.3  Legal", level=2, color="7B2D8B")
    table(doc,
        ["Scenario", "Without TrustLayer", "With TrustLayer"],
        [
            ("Paralegal asks for case law on non-competes in California",
             "AI cites a fabricated case: 'Smith v. TechCorp (9th Cir. 2022)'. Case does not exist. Filed in a motion. Attorney sanctioned.",
             "Source Verification detects unverifiable citation. Response FLAGGED with warning: 'Verify in Westlaw before use.'"),
            ("Lawyer asks about a specific statute",
             "AI cites a statute number that does not exist in the relevant jurisdiction.",
             "Pattern Recognition and Source Verification detect fabricated statute. BLOCKED."),
        ],
        "7B2D8B",
        col_widths=[1.8, 2.4, 2.6]
    )

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 6 — WHAT THE USER ACTUALLY SEES
    # ══════════════════════════════════════════════════════════
    heading(doc, "6.  What the End User Actually Sees", color="1F3864")
    para(doc,
        "From the end user's perspective, TrustLayer is invisible when everything works correctly. "
        "They simply receive accurate answers. TrustLayer only becomes visible when it intervenes."
    )
    doc.add_paragraph()

    table(doc,
        ["Decision", "What the user sees", "What actually happened behind the scenes"],
        [
            ("PASS",
             "The AI's response appears normally, instantly. No delay visible to user.",
             "TrustLayer scored the response at 75%+ confidence. All 8 checks passed. Grounding verified. Response delivered."),
            ("FLAG",
             "User may see: 'A specialist is reviewing your question and will respond shortly.' OR response is shown with a disclaimer banner.",
             "TrustLayer scored 50-74%. A compliance reviewer sees the flagged response in a review queue with the detection breakdown."),
            ("BLOCK",
             "User sees a safe fallback: 'I am not able to provide that specific information. Please speak with one of our advisors.'",
             "TrustLayer scored below 50% or detected a prohibited claim. The original AI response is logged but never shown to the user."),
        ],
        "1F3864",
        col_widths=[1.0, 2.5, 3.3]
    )

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 7 — THE ANALOGY BANK
    # ══════════════════════════════════════════════════════════
    heading(doc, "7.  The Analogy Bank  (use whichever lands best)", color="1F3864")
    para(doc,
        "When explaining TrustLayer in conversation, one of these analogies will resonate "
        "depending on your listener's background."
    )
    doc.add_paragraph()

    analogies = [
        ("The Seatbelt",
         "An airbag does not make driving slower — it makes it safer. "
         "TrustLayer does not slow down your AI product — it makes it safe to deploy. "
         "You would not sell a car without a seatbelt. Do not deploy an enterprise AI without TrustLayer."),
        ("The Spell Checker, but for Facts",
         "Microsoft Word does not stop you from writing — it underlines mistakes. "
         "TrustLayer does not stop your AI from generating responses — it flags factual mistakes "
         "before anyone reads them. The difference is that a spell checker catches typos; "
         "TrustLayer catches claims that could cost your company millions."),
        ("The Compliance Officer at the Door",
         "Imagine every answer your AI gives has to pass through a compliance officer before "
         "reaching the customer. The officer has a copy of your product manual, your rate sheet, "
         "and your legal guidelines. They read every response in 2 seconds and either wave it "
         "through, send it for review, or replace it with a safe alternative. That is TrustLayer."),
        ("The ATM Cash Verification",
         "When you deposit money at an ATM, the machine counts the bills before crediting your account. "
         "It does not trust the cash blindly — it verifies it. TrustLayer is the verification step "
         "your AI stack is currently missing."),
        ("Air Traffic Control",
         "Pilots are highly trained and usually correct. But air traffic control exists because "
         "'usually correct' is not good enough when lives are at stake. "
         "AI models are impressive and often right. But in banking and healthcare, 'often right' "
         "is not good enough. TrustLayer is air traffic control for your AI."),
    ]

    for title, body in analogies:
        p = doc.add_paragraph()
        r = p.add_run(title + ":  ")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = rgb("2E75B6")
        r2 = p.add_run(body)
        r2.font.size = Pt(11)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 8 — BUSINESS VALUE
    # ══════════════════════════════════════════════════════════
    heading(doc, "8.  Business Value — The Numbers Argument", color="1F3864")
    para(doc,
        "When speaking to a C-suite audience or a buyer, frame TrustLayer in terms of "
        "cost avoidance and risk reduction, not technology."
    )
    doc.add_paragraph()

    table(doc,
        ["Risk Category", "Cost Without TrustLayer", "TrustLayer Prevention"],
        [
            ("Regulatory fine (BFSI)",
             "CFPB fines for misleading AI advice: $1M to $100M+",
             "BLOCK on illegal guarantees and fabricated rates before customer exposure"),
            ("Patient safety (Healthcare)",
             "Wrongful harm lawsuit for AI-recommended wrong dosage: $5M to $50M+",
             "BLOCK on dosages deviating from FDA grounding data"),
            ("Legal sanctions",
             "Attorney sanctions and case dismissal for fabricated citations: career-ending",
             "FLAG on unverifiable case citations before filing"),
            ("Brand damage",
             "Viral social media post: 'Bank AI promised me 4.2% mortgage' — immeasurable",
             "BLOCK stops the false claim before the customer screenshot is taken"),
            ("Compliance audit failure",
             "EU AI Act non-compliance: up to 3% of global annual revenue",
             "Built-in audit trail, human oversight, risk classification — 94% EU AI Act coverage"),
        ],
        "1F3864",
        col_widths=[1.7, 2.5, 2.6]
    )

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 9 — MINIMUM VIABLE INTEGRATION
    # ══════════════════════════════════════════════════════════
    heading(doc, "9.  Minimum Viable Integration  (3 lines of code)", color="1F3864")
    para(doc,
        "For a technical audience, show them how simple the integration actually is. "
        "TrustLayer is not a rip-and-replace. It is an additive layer."
    )
    doc.add_paragraph()

    heading(doc, "Before TrustLayer:", level=2, color="CC2222")
    code_before = doc.add_paragraph()
    code_before.paragraph_format.left_indent = Inches(0.4)
    r = code_before.add_run(
        "response = claude_client.messages.create(...)  # raw response\n"
        "return response.content[0].text                # delivered to user with no verification"
    )
    r.font.name = "Courier New"
    r.font.size = Pt(10)
    r.font.color.rgb = rgb("CC2222")

    doc.add_paragraph()
    heading(doc, "After TrustLayer (3 lines added):", level=2, color="007A3D")
    code_after = doc.add_paragraph()
    code_after.paragraph_format.left_indent = Inches(0.4)
    r2 = code_after.add_run(
        "response = claude_client.messages.create(...)\n"
        "ai_text   = response.content[0].text\n"
        "\n"
        "result = detector.analyze(query, ai_text, system_prompt, grounding_ctx)\n"
        "\n"
        "if result.decision == 'BLOCK':\n"
        "    return SAFE_FALLBACK_MESSAGE\n"
        "elif result.decision == 'FLAG':\n"
        "    route_to_human_review(result)\n"
        "else:\n"
        "    return ai_text  # PASS — delivered to user"
    )
    r2.font.name = "Courier New"
    r2.font.size = Pt(10)
    r2.font.color.rgb = rgb("007A3D")

    doc.add_paragraph()
    para(doc,
        "That is the complete integration. No model swap. No app rebuild. No infrastructure change. "
        "Three lines of code added to the existing response handler, and every AI output is now "
        "governed before delivery."
    )

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 10 — ELEVATOR PITCHES BY TIME
    # ══════════════════════════════════════════════════════════
    heading(doc, "10.  Elevator Pitches by Time Available", color="1F3864")

    pitches = [
        ("30 seconds",
         "Your AI chatbot will eventually give a customer a wrong interest rate, a fake promotion, "
         "or an impossible guarantee. When that happens without TrustLayer, the customer reads it "
         "and you have a complaint, a fine, or a lawsuit. With TrustLayer, that response never "
         "reaches the customer — it is stopped, logged, and replaced with a safe message "
         "before anyone sees it."),
        ("2 minutes",
         "We built a real-time hallucination detection layer called TrustLayer that sits between "
         "your AI assistant and your customers. Every AI response passes through 8 independent "
         "checks in under 2 seconds. We verify every specific claim against your actual product "
         "data — current rates, eligibility rules, policy limits. If the AI fabricated a rate or "
         "invented a promotion, TrustLayer catches it before the customer reads it and either "
         "blocks the response or routes it to a human reviewer.\n\n"
         "We also cross-validate every response using a second independent AI model. "
         "When the two models disagree, we treat that disagreement itself as a risk signal "
         "and escalate the decision. The end result: a PASS, FLAG, or BLOCK decision on "
         "every single response, with a full audit trail for compliance."),
        ("5 minutes",
         "Use the full live demo. Run the recommended script from Section 4 of the "
         "Banking Demo Questions document: start with a safe question that passes, "
         "then show a hallucination being flagged, then show a block on a fabricated rate. "
         "End with: 'That just happened in real time. The AI generated a wrong answer. "
         "TrustLayer caught it. The customer never saw it.'"),
    ]

    for time_label, pitch_text in pitches:
        p = doc.add_paragraph()
        r = p.add_run(time_label + " pitch:")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = rgb("1F3864")
        body = doc.add_paragraph()
        body.paragraph_format.left_indent = Inches(0.3)
        body.add_run(pitch_text).font.size = Pt(11)
        doc.add_paragraph()

    # ── Footer ──────────────────────────────────────────────────
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(
        "TrustLayer AI  —  Detect · Prevent · Govern  |  Canada Hackathon 2026"
    )
    fr.italic = True
    fr.font.size = Pt(10)
    fr.font.color.rgb = rgb("888888")

    out = r"C:\aiprojects\aitrust_prod\TrustLayer_AI_RealWorld_Deployment_Guide.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
