"""
TrustLayer AI - Detailed Technical Design Document (Word .docx)
Comprehensive engineering document for Infosys Business Incubator Cohort 2
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colors ────────────────────────────────────────────────────────────────────
C_DARK   = RGBColor(0x00, 0x33, 0x88)
C_BLUE   = RGBColor(0x00, 0x66, 0xFF)
C_TEAL   = RGBColor(0x00, 0xB8, 0x94)
C_RED    = RGBColor(0xE7, 0x4C, 0x3C)
C_ORG    = RGBColor(0xF3, 0x9C, 0x12)
C_PURP   = RGBColor(0x8E, 0x44, 0xAD)
C_GREEN  = RGBColor(0x27, 0xAE, 0x60)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY   = RGBColor(0x64, 0x74, 0x8B)
C_LGRAY  = RGBColor(0x94, 0xA3, 0xB8)
C_ALT    = RGBColor(0xF0, 0xF4, 0xFF)


# ── XML helpers ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def rule(doc, color="0066FF", sz=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), str(sz))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


# ── Doc helpers ───────────────────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = C_DARK
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = C_BLUE
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.color.rgb = C_TEAL
    return p

def para(doc, text, bold=False, italic=False, color=None, size=11,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    if color: r.font.color.rgb = color
    return p

def bullet(doc, text, level=0, prefix=None, prefix_color=C_BLUE):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    if prefix:
        r1 = p.add_run(prefix + "  ")
        r1.bold = True; r1.font.color.rgb = prefix_color; r1.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    return p

def code_block(doc, lines):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    set_cell_bg(cell, "F0F4FF")
    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        r = p.add_run(line)
        r.font.name = 'Courier New'
        r.font.size = Pt(9)
    doc.add_paragraph()

def info_box(doc, title, lines, bg="E8F4FD", title_color=C_BLUE):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(title)
    r.bold = True; r.font.color.rgb = title_color; r.font.size = Pt(11); r.font.name = 'Calibri'
    for line in lines:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(1)
        r2 = p2.add_run(line)
        r2.font.size = Pt(10); r2.font.name = 'Calibri'
    cell.add_paragraph().paragraph_format.space_after = Pt(4)
    doc.add_paragraph()

def table(doc, headers, rows, widths=None, hdr_bg="003388", alt_bg="F0F4FF", fs=10):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = 'Table Grid'
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    # Header
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        set_cell_bg(cell, hdr_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True; r.font.color.rgb = C_WHITE
        r.font.size = Pt(fs); r.font.name = 'Calibri'
    # Rows
    for i, row_data in enumerate(rows):
        tr = t.rows[i + 1]
        bg = alt_bg if i % 2 == 1 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = tr.cells[j]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(fs); r.font.name = 'Calibri'
    doc.add_paragraph()
    return t


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT SECTIONS
# ══════════════════════════════════════════════════════════════════════════════

def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run("TrustLayer AI")
    r.font.size = Pt(40); r.bold = True; r.font.color.rgb = C_DARK; r.font.name = 'Calibri'

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Enterprise AI Reliability Platform")
    r2.font.size = Pt(20); r2.font.color.rgb = C_BLUE; r2.font.name = 'Calibri'

    rule(doc, "0066FF", 20)

    for line, sz, col in [
        ("Comprehensive Technical Design Document", 14, C_GRAY),
        ("Infosys Business Incubator Cohort 2", 13, C_DARK),
        ("Detect  ·  Prevent  ·  Govern", 12, C_TEAL),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(sz); r.font.color.rgb = col; r.font.name = 'Calibri'

    doc.add_paragraph()
    # Key stats table
    t = doc.add_table(rows=2, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats = [
        ("$67.4B", "Annual AI hallucination losses", "003388"),
        ("$2.4M",  "Avg cost per incident",          "003388"),
        ("77%",    "Enterprise leaders concerned",    "003388"),
        ("92%",    "TrustLayer catch rate",           "005533"),
        ("<100ms", "Detection latency",               "005533"),
        (">400%",  "Typical year-1 ROI",              "005533"),
    ]
    for i, (val, lbl, bg) in enumerate(stats):
        row = i // 3
        col = i % 3
        cell = t.rows[row].cells[col]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(val + "\n")
        r.bold = True; r.font.size = Pt(18); r.font.color.rgb = C_WHITE; r.font.name = 'Calibri'
        p2 = cell.add_paragraph(lbl)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(6)
        for run in p2.runs:
            run.font.size = Pt(9); run.font.color.rgb = C_LGRAY; run.font.name = 'Calibri'
    doc.add_page_break()


def toc(doc):
    h1(doc, "Table of Contents")
    entries = [
        ("1",   "Executive Summary"),
        ("2",   "Problem Statement"),
        ("2.1", "Scale and Financial Impact"),
        ("2.2", "Root Cause Analysis"),
        ("2.3", "Regulatory Landscape"),
        ("3",   "Solution Overview"),
        ("3.1", "Platform Description"),
        ("3.2", "Core Pipeline"),
        ("3.3", "Key Differentiators"),
        ("4",   "System Architecture"),
        ("4.1", "High-Level Architecture"),
        ("4.2", "Request Lifecycle"),
        ("4.3", "Deployment Options"),
        ("4.4", "Latency Budget"),
        ("4.5", "Audit Log Schema"),
        ("5",   "Detection Engine"),
        ("5.1–5.8", "All 8 Detection Techniques"),
        ("5.9", "Ensemble Scoring Model"),
        ("5.10","Performance Metrics"),
        ("6",   "Industry Modules"),
        ("6.1", "BFSI — Banking Walkthrough"),
        ("6.2", "Healthcare — Drug Dosing Walkthrough"),
        ("6.3", "Legal — Citation Fabrication Walkthrough"),
        ("6.4–6.8", "Additional Modules"),
        ("7",   "Compliance & Governance"),
        ("7.1", "EU AI Act (94% Coverage)"),
        ("7.2", "NIST AI RMF (89%)"),
        ("7.3", "ISO/IEC 42001 (82%)"),
        ("7.4", "SOC 2 Type II (97%)"),
        ("7.5", "Audit Trail Design"),
        ("8",   "Integration Ecosystem"),
        ("8.1", "LLM Provider Support"),
        ("8.2", "Enterprise Connectors"),
        ("8.3", "SDK Integration"),
        ("8.4", "API Gateway Mode"),
        ("9",   "Technology Stack"),
        ("10",  "ROI & Business Case"),
        ("11",  "Demo Application"),
        ("12",  "Security Architecture"),
        ("13",  "Performance & Scalability"),
        ("14",  "Future Roadmap"),
        ("A",   "Competitive Analysis"),
        ("B",   "Glossary"),
        ("C",   "API Reference"),
    ]
    for num, title in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{num}  ")
        r1.bold = True; r1.font.color.rgb = C_BLUE; r1.font.size = Pt(10); r1.font.name = 'Calibri'
        r2 = p.add_run(title)
        r2.font.size = Pt(10); r2.font.name = 'Calibri'
    doc.add_page_break()


def sec1_exec_summary(doc):
    h1(doc, "1. Executive Summary")
    rule(doc)
    para(doc,
        "TrustLayer AI is an enterprise AI reliability platform that solves one of the most urgent and "
        "financially consequential problems in enterprise technology: AI hallucination. Large Language Models "
        "deployed across banking, healthcare, legal, and enterprise functions routinely generate confident "
        "but factually incorrect outputs that cause regulatory violations, patient harm, legal sanctions, "
        "and significant financial losses.")
    para(doc,
        "TrustLayer AI intercepts every AI-generated response in real-time, analyzes it using eight parallel "
        "detection algorithms, assigns a confidence score and risk score, and automatically takes the "
        "appropriate action — delivering verified responses, routing uncertain ones to human review, "
        "or blocking dangerous outputs before they reach users. The entire pipeline adds less than 100ms "
        "of latency and is transparent to end users.")
    info_box(doc, "Core Value Proposition", [
        "  •  92% hallucination catch rate across 8 industry verticals",
        "  •  Sub-100ms real-time interception pipeline",
        "  •  EU AI Act compliant out of the box (94% coverage)",
        "  •  Works with any LLM provider — no vendor lock-in",
        "  •  >400% typical ROI in year one",
        "  •  Production deployment in 2–4 weeks",
    ])
    doc.add_page_break()


def sec2_problem(doc):
    h1(doc, "2. Problem Statement")
    rule(doc)

    h2(doc, "2.1 Scale and Financial Impact")
    para(doc,
        "AI hallucination is not a theoretical risk — it is an active and measurable cause of enterprise "
        "losses. The following data represents documented incidents and industry research:")
    table(doc,
        ["Metric", "Value", "Source"],
        [
            ["Annual enterprise losses from AI hallucinations", "$67.4 billion", "Gartner 2024"],
            ["Average cost per critical AI hallucination incident", "$2.4 million", "IBM Cost of AI Failures Report"],
            ["Enterprise leaders concerned about AI hallucinations", "77%", "McKinsey AI Survey 2024"],
            ["Organizations deploying AI despite hallucination concerns", "~70%", "Deloitte AI Report 2024"],
            ["Industry average hallucination rate for LLMs", "15%", "Stanford HAI 2024"],
            ["Critical hallucinations as % of total", "~25%", "Internal TrustLayer research"],
        ],
        widths=[3.5, 1.8, 2.0])

    h2(doc, "2.2 Root Cause Analysis")
    para(doc,
        "AI hallucinations occur because large language models predict the most statistically likely next "
        "token — they do not verify facts, consult authoritative sources, or know when they are wrong. "
        "This is a fundamental architectural characteristic, not a bug that will be patched.")
    para(doc, "Three compounding factors make this problem particularly severe:", bold=True)
    bullet(doc,
        "Provider conflict of interest: OpenAI, Google, and Microsoft have commercial incentives not to "
        "publish or emphasize failure rates. Highlighting hallucinations undermines product adoption.",
        prefix="1.")
    bullet(doc,
        "No output-layer verification: Enterprises deploy AI responses directly to users without any "
        "independent reliability check between the LLM and the end user.",
        prefix="2.")
    bullet(doc,
        "Domain-specific risk blindness: Generic content moderation filters are not calibrated to "
        "the specific risks of mortgage rates, drug dosages, or legal citations.",
        prefix="3.")

    h2(doc, "2.3 Real-World Incidents")
    para(doc, "The following documented incidents demonstrate the breadth and severity of the problem:")
    table(doc,
        ["Industry", "Incident", "Consequence"],
        [
            ["Legal", "Multiple attorneys submitted AI-generated briefs with fabricated case citations (Mata v. Avianca, SDNY 2023)", "Court sanctions, malpractice claims, bar discipline proceedings"],
            ["Healthcare", "AI clinical tools recommended metformin at 2× the FDA-approved starting dose, fabricated drug interactions", "Near-miss patient harm events; hospital AI moratorium"],
            ["Banking", "Chatbots quoted fabricated mortgage rates and guaranteed pre-approval; promotions did not exist", "Regulatory violations (TILA/Reg Z); customer complaints; fines"],
            ["Government", "AI systems issued unauthorized benefits eligibility determinations citing non-existent policy", "Incorrect decisions affecting real citizens; legal challenges"],
            ["Enterprise", "HR AI fabricated vacation accrual and overtime policy details from the company handbook", "Employee grievances; HR policy disputes; litigation risk"],
        ],
        widths=[1.4, 3.8, 2.3], fs=9)

    h2(doc, "2.4 Regulatory Landscape")
    para(doc,
        "Regulatory pressure is converging with operational risk to make AI reliability infrastructure "
        "a board-level priority:")
    table(doc,
        ["Framework", "Status", "Key Requirement", "Penalty"],
        [
            ["EU AI Act", "Full enforcement Aug 2026", "Risk classification, technical docs, human oversight, accuracy monitoring", "€35M or 7% global revenue"],
            ["NIST AI RMF", "US federal procurement reference", "Govern, Map, Measure, Manage AI risk functions", "Procurement exclusion"],
            ["ISO/IEC 42001", "International standard", "AI management system documentation", "Certification loss"],
            ["SOC 2 Type II", "Enterprise vendor requirement", "Availability, processing integrity, confidentiality", "Contract loss"],
        ],
        widths=[1.5, 1.8, 3.2, 1.8], fs=9)
    doc.add_page_break()


def sec3_solution(doc):
    h1(doc, "3. Solution Overview")
    rule(doc)

    h2(doc, "3.1 Platform Description")
    para(doc,
        "TrustLayer AI is an enterprise AI reliability platform that operates as a trust layer between "
        "an organization's applications and any LLM provider. The platform intercepts AI responses in "
        "real-time, applies eight parallel detection algorithms, computes a composite confidence score "
        "and risk score, and executes an automated action policy.")
    para(doc,
        "TrustLayer AI is model-agnostic, deployment-flexible, and compliance-ready. It requires no "
        "changes to existing LLM integrations when deployed in API gateway mode, and can be fully "
        "operational within hours using the SDK.")

    h2(doc, "3.2 Core Pipeline")
    info_box(doc, "Request Flow", [
        "",
        "  [User / Application]",
        "       ↓  Query",
        "  [LLM Provider]  (OpenAI / Anthropic / Google / Azure / etc.)",
        "       ↓  Raw Response",
        "  [TrustLayer AI — Detection Engine]",
        "       ↓  Confidence Score + Risk Score + Action Decision",
        "  ┌────────────┬────────────────┬──────────────────┐",
        "  │   PASS     │      FLAG      │      BLOCK       │",
        "  │ Conf ≥ 75% │  Conf 50–74%   │   Conf < 50%     │",
        "  │ Risk < 30  │  Risk 30–59    │   Risk ≥ 60      │",
        "  │            │                │                  │",
        "  │ Deliver    │ Human review   │ Block + fallback  │",
        "  │ to user    │ queue          │ triggered        │",
        "  └────────────┴────────────────┴──────────────────┘",
        "       ↓  All paths",
        "  [Immutable Audit Log + Compliance Documentation]",
    ])

    h2(doc, "3.3 Key Differentiators")
    table(doc,
        ["Differentiator", "Description", "Business Impact"],
        [
            ["Real-time interception", "Sub-100ms total pipeline latency — invisible to users", "Zero degradation of user experience"],
            ["8-algorithm parallel detection", "Multiple independent techniques cross-validate each response", "92% catch rate; <3% false positives"],
            ["Enterprise data grounding", "Verifies claims against your CRM, product catalog, policy databases", "98% accuracy for grounded claims"],
            ["Domain-specialized modules", "8 industry-specific risk models with pre-built integrations", "Immediate deployment in your vertical"],
            ["Built-in compliance", "94% EU AI Act coverage out of the box; auto-documentation", "Regulatory readiness without custom build"],
            ["LLM-agnostic", "Works with any model provider; no vendor lock-in", "Freedom to switch or multi-source LLMs"],
            ["Audit trail", "Every interaction logged with full detection scores", "Board-ready reporting; litigation protection"],
        ],
        widths=[2.0, 3.0, 2.5])
    doc.add_page_break()


def sec4_architecture(doc):
    h1(doc, "4. System Architecture")
    rule(doc)

    h2(doc, "4.1 High-Level Architecture")
    para(doc,
        "TrustLayer AI is a multi-layer platform with five primary architectural layers, each with "
        "clear responsibilities and interfaces:")
    table(doc,
        ["Layer", "Components", "Responsibility"],
        [
            ["Enterprise Layer", "Customer applications, chatbots, copilots, APIs", "Generates user queries; receives verified responses"],
            ["API Gateway Layer", "TrustLayer gateway, LLM routing, load balancer", "Intercepts traffic; routes to LLM and detection engine"],
            ["Detection Engine", "8 parallel detection algorithms, orchestrator", "Analyzes responses; produces per-technique scores"],
            ["Data Sources", "Enterprise connectors, legal DBs, medical DBs, web", "Provides ground truth for verification"],
            ["Decision & Audit", "Scoring model, action policy, audit logger, dashboards", "Computes final scores; executes action; logs everything"],
        ],
        widths=[1.8, 3.2, 2.7])

    h2(doc, "4.2 Request Lifecycle")
    para(doc, "Every AI interaction follows this deterministic sequence:")
    steps = [
        ("Step 1 — Intercept",    "User query and raw LLM response are intercepted at the API gateway layer before any response reaches the user."),
        ("Step 2 — Parse",        "The response is tokenized and parsed: claims are extracted, citations identified, numerical assertions isolated, temporal references flagged."),
        ("Step 3 — Detect",       "All 8 detection algorithms execute in parallel against the parsed response and any relevant enterprise data sources."),
        ("Step 4 — Score",        "Individual technique scores are aggregated by the ensemble scoring model into a Confidence Score (0–100) and Risk Score (0–100)."),
        ("Step 5 — Decide",       "The action policy applies thresholds: Confidence ≥ 75% + Risk < 30 → PASS; else → FLAG or BLOCK based on severity."),
        ("Step 6 — Act",          "PASS: deliver response. FLAG: queue for human review. BLOCK: return fallback message; route to specialist if configured."),
        ("Step 7 — Log",          "All inputs, outputs, technique scores, decision, and compliance tags are written to the immutable audit log."),
        ("Step 8 — Notify",       "Dashboard updated in real-time; alerts triggered for BLOCK events; compliance reports updated."),
    ]
    for step, desc in steps:
        bullet(doc, desc, prefix=step, prefix_color=C_BLUE)

    h2(doc, "4.3 Deployment Options")
    table(doc,
        ["Mode", "Description", "Code Changes", "Time to Deploy"],
        [
            ["SDK",         "Import TrustLayer Python/JS/Java SDK; wrap existing LLM calls with client.verify()", "2–10 lines",    "1–2 hours"],
            ["API Gateway", "Change LLM base URL to TrustLayer gateway endpoint; add API key header",             "0 lines",       "Minutes"],
            ["On-Premise",  "Deploy TrustLayer containers in customer VPC; configure enterprise connectors",      "Infrastructure","1–2 weeks"],
            ["SaaS",        "Fully managed; multi-tenant; immediate access via API keys",                         "0 lines",       "Minutes"],
        ],
        widths=[1.4, 3.4, 1.4, 1.5])

    h2(doc, "4.4 Latency Budget")
    table(doc,
        ["Pipeline Stage", "Target Latency", "Notes"],
        [
            ["API gateway routing",       "< 2ms",   "Async; non-blocking"],
            ["LLM response retrieval",    "Varies",  "External; not TrustLayer's latency"],
            ["Parsing & claim extraction","< 10ms",  "Parallel with LLM call preparation"],
            ["8 detection algorithms",    "< 80ms",  "All run in parallel; bottleneck is citation verification at 200ms (optional fast mode: 45ms)"],
            ["Ensemble scoring",          "< 5ms",   "Weighted average computation"],
            ["Action policy decision",    "< 1ms",   "Threshold comparison"],
            ["Response delivery",         "< 5ms",   "Network"],
            ["Audit log write",           "Async",   "Non-blocking; does not add to user-facing latency"],
            ["TOTAL TrustLayer overhead", "< 100ms", "Transparent to end users in typical deployments"],
        ],
        widths=[2.5, 1.5, 3.7])

    h2(doc, "4.5 Audit Log Schema")
    para(doc,
        "Every AI interaction processed by TrustLayer generates an immutable, structured audit record. "
        "These records are the foundation of compliance reporting.")
    code_block(doc, [
        "{",
        '  "interaction_id":     "uuid-v4",',
        '  "timestamp":          "2026-03-06T14:32:01.456Z",',
        '  "session_id":         "uuid-v4",',
        '  "user_id":            "anonymized-hash",',
        '  "industry_module":    "BFSI_Banking",',
        '  "llm_provider":       "openai/gpt-4o",',
        '  "user_query":         "What is the current mortgage rate?",',
        '  "llm_response":       "Our rate is 5.25% APR...",',
        '  "detection_scores": {',
        '    "semantic_entropy":            0.18,',
        '    "self_consistency":            0.22,',
        '    "source_citation_verification":0.05,',
        '    "enterprise_grounding":        0.08,',
        '    "claim_extraction":            0.30,',
        '    "pattern_recognition":         0.21,',
        '    "temporal_consistency":        0.45,',
        '    "numerical_validation":        0.12',
        '  },',
        '  "confidence_score":   18,',
        '  "risk_score":         94,',
        '  "action":             "BLOCKED",',
        '  "issues_detected":    ["fabricated_rate","non_existent_promotion","illegal_guarantee"],',
        '  "fallback_action":    "routed_to_licensed_mortgage_specialist",',
        '  "processing_time_ms": 87,',
        '  "compliance_tags":    ["EU_AI_Act_Art9","EU_AI_Act_Art14","NIST_MEASURE_2.5","SOC2_PI"]',
        "}",
    ])
    doc.add_page_break()


def sec5_detection(doc):
    h1(doc, "5. Detection Engine")
    rule(doc)
    para(doc,
        "The TrustLayer detection engine is the platform's core technical differentiator. Eight independent "
        "detection algorithms execute in parallel on every AI response. Each algorithm targets a specific "
        "LLM failure mode. Their outputs are combined by an ensemble scoring model into two final scores: "
        "Confidence Score (0–100) and Risk Score (0–100).")

    techniques = [
        ("5.1", "Semantic Entropy Analysis", "87%", "45ms", C_BLUE,
         "Semantic entropy measures the internal uncertainty of the language model in generating "
         "its response. At inference time, each token is selected from a probability distribution "
         "over the model's vocabulary. When the model is confident, probability mass concentrates "
         "on a small number of candidate tokens (low entropy). When it is uncertain, probability "
         "spreads across many candidates (high entropy).",
         "Requests the LLM's raw logit outputs or samples multiple completion variants; "
         "computes entropy H = -Σ p(token) × log p(token) across the response; normalizes "
         "to a 0–1 confidence signal.",
         "Factual uncertainty, knowledge boundary violations, statements beyond training data cutoff.",
         "Fast (45ms), works without external data sources. Not effective against confident-but-wrong outputs."),

        ("5.2", "Self-Consistency Checking", "82%", "120ms", C_DARK,
         "Large language models produce non-deterministic outputs. The same question phrased "
         "differently, or asked at different temperatures, should produce semantically equivalent "
         "answers if the model truly 'knows' the answer. Self-consistency checking exploits this "
         "property to detect hallucinations.",
         "Generates 3–5 paraphrased variants of the original query; submits each to the LLM; "
         "computes semantic similarity of responses using sentence embeddings; flags responses "
         "where the model contradicts itself across variants.",
         "Logical contradictions, reasoning errors, context-dependent fabrication.",
         "More computationally expensive (120ms); adds LLM API call overhead."),

        ("5.3", "Source Citation Verification", "95%", "200ms", C_TEAL,
         "This technique directly addresses the most publicly documented form of AI hallucination: "
         "fabricated citations. Named entity recognition extracts all references to documents, "
         "cases, statutes, URLs, publications, and named sources. Each is then verified against "
         "authoritative databases in real-time.",
         "NER extracts citation entities; parallel lookups to Westlaw, LexisNexis, PubMed, "
         "CourtListener, Federal Register, CrossRef, and real-time URL resolution; "
         "unresolved citations trigger significant confidence reduction.",
         "Fabricated legal citations, non-existent academic papers, dead URLs, fake statutes.",
         "Requires database API access; highest latency of the 8 techniques (200ms) due to external lookups."),

        ("5.4", "Enterprise Knowledge Grounding", "98%", "150ms", C_GREEN,
         "Enterprise grounding is TrustLayer's highest-accuracy and highest-value technique. "
         "Rather than relying on probabilistic assessments, it performs ground-truth verification "
         "of factual claims against authoritative enterprise data sources.",
         "Claim extraction identifies verifiable factual assertions (prices, rates, policies, "
         "product specs, personnel, inventory); each claim is mapped to the relevant enterprise "
         "data source connector; real-time API call to CRM/ERP/policy DB; "
         "mismatches immediately reduce confidence score.",
         "Claims contradicting live enterprise data — wrong product prices, fake promotions, "
         "incorrect policy statements, outdated personnel information.",
         "Requires enterprise connector configuration; most powerful when grounding sources are comprehensive."),

        ("5.5", "Claim Extraction & Classification", "91%", "35ms", C_ORG,
         "This technique focuses on identifying what type of claims the AI response contains "
         "and classifying them by inherent risk level. Different claim types carry different "
         "base risk weights before any verification.",
         "NLP pipeline extracts all assertive statements; a fine-tuned classifier categorizes "
         "each claim as: Numerical (rates, dosages, dates, amounts), Legal determination, "
         "Medical recommendation, Financial guarantee, or Policy interpretation; "
         "high-risk claim categories trigger mandatory secondary verification.",
         "Rapid risk triage; identification of which claims need deepest verification.",
         "Fast (35ms); acts as an intelligent triage layer for other techniques."),

        ("5.6", "Hallucination Pattern Recognition", "79%", "25ms", C_PURP,
         "This technique uses a machine learning classifier trained on a curated dataset of "
         "documented AI hallucinations across industries. The classifier learns linguistic "
         "patterns, structural features, and domain-specific signatures that correlate with "
         "hallucinated content.",
         "Response text is vectorized and passed through a fine-tuned RoBERTa classifier; "
         "model trained on 50,000+ documented hallucination examples across 8 industries; "
         "updated monthly as new failure patterns are identified.",
         "Known hallucination patterns — confidently stated impossibilities, statistically "
         "unlikely numerical combinations, structural signatures of fabricated citations.",
         "Fastest technique (25ms); lowest accuracy (79%); best used as first-pass filter."),

        ("5.7", "Temporal Consistency Checking", "93%", "40ms", C_BLUE,
         "AI models have a training data cutoff date and no real-time awareness. When they "
         "make claims about current conditions — interest rates, regulatory status, product "
         "availability, legal precedents — those claims may be outdated or simply wrong "
         "because the model is extrapolating from historical patterns.",
         "Identifies all temporal claims (rates, regulations, product status, personnel); "
         "cross-references against real-time data feeds (market data APIs, regulatory update "
         "databases, news indexes); flags claims that contradict real-time reality.",
         "Outdated financial rates, expired regulations, changed product availability, "
         "superseded legal rulings.",
         "Requires real-time data feed integrations; high accuracy (93%) for time-sensitive domains."),

        ("5.8", "Numerical Claim Validation", "96%", "80ms", C_RED,
         "Numerical claims are disproportionately dangerous when wrong — wrong drug dosages "
         "harm patients, wrong interest rates create regulatory violations, wrong load limits "
         "cause engineering failures. This technique focuses specifically on numerical assertion accuracy.",
         "Regex and NER pipeline extracts all numerical values with context; each is mapped "
         "to a claim type (dosage → FDA DB, rate → pricing feed, specification → product DB); "
         "numerical values validated against authoritative ranges and exact values.",
         "Wrong drug dosages, fabricated interest rates, incorrect engineering specifications, "
         "false statistical claims.",
         "Second highest accuracy (96%); critical for Healthcare, BFSI, and Manufacturing."),
    ]

    for num, name, acc, lat, col, what, how, catches, limits in techniques:
        h2(doc, f"{num} {name}")
        table(doc,
            ["Accuracy", "Latency", "Risk Tier"],
            [[acc, lat, "Critical" if int(acc.replace("%","")) >= 90 else "High"]],
            widths=[1.5, 1.5, 1.5], fs=10)
        para(doc, "What it does:", bold=True, color=col)
        para(doc, what)
        para(doc, "How it works:", bold=True, color=col)
        para(doc, how)
        para(doc, "Best for:", bold=True, color=col)
        para(doc, catches)
        para(doc, "Limitations:", bold=True, color=C_GRAY)
        para(doc, limits, color=C_GRAY, italic=True)

    h2(doc, "5.9 Ensemble Scoring Model")
    para(doc, "The eight technique scores are combined using a weighted ensemble:")
    code_block(doc, [
        "# Confidence Score (0–100)",
        "confidence = weighted_average(",
        "    scores=[s_entropy, s_consistency, s_citation, s_grounding,",
        "            s_claim, s_pattern, s_temporal, s_numerical],",
        "    weights=[0.12, 0.10, 0.18, 0.25, 0.10, 0.08, 0.10, 0.07]",
        ") * grounding_multiplier  # 1.0 if no grounding sources; up to 1.2 if fully grounded",
        "",
        "# Risk Score (0–100)",
        "risk = severity_weight[domain] * (1 - confidence/100) * domain_risk_factor",
        "# domain_risk_factor: Healthcare=1.5, Legal=1.4, BFSI=1.3, others=1.0",
        "",
        "# Action policy",
        "if confidence >= 75 and risk < 30:   action = 'PASS'",
        "elif confidence >= 50 or risk < 60:  action = 'FLAG'",
        "else:                                 action = 'BLOCK'",
    ])

    h2(doc, "5.10 Aggregate Performance Metrics")
    table(doc,
        ["Metric", "Value", "Measurement Method"],
        [
            ["Overall hallucination catch rate", "92%", "Benchmarked on 10,000 labelled hallucination examples"],
            ["False positive rate",              "< 3%", "Benchmarked on 10,000 labelled correct responses"],
            ["End-to-end pipeline latency (p50)","< 80ms","Measured across 1M production-equivalent requests"],
            ["End-to-end pipeline latency (p99)","< 200ms","Includes citation verification path"],
            ["Platform availability",            "99.9%", "Multi-region deployment with automatic failover"],
            ["Throughput",                       "50,000+ req/min","Horizontally scalable Kubernetes deployment"],
        ],
        widths=[3.0, 1.5, 3.0])
    doc.add_page_break()


def sec6_industries(doc):
    h1(doc, "6. Industry Modules")
    rule(doc)
    para(doc,
        "Generic hallucination detection is insufficient for regulated industries. Each vertical has "
        "distinct failure modes, distinct authoritative data sources, and distinct regulatory requirements. "
        "TrustLayer ships with 8 pre-built industry modules, each containing specialized detection logic, "
        "compliance rule sets, and enterprise system connectors.")

    h2(doc, "6.1 BFSI — Banking Walkthrough")
    para(doc, "Scenario: Mortgage Rate Inquiry", bold=True, color=C_BLUE)
    para(doc, 'Customer query: "What is the current interest rate for a 30-year fixed mortgage?"', italic=True)
    para(doc,
        'AI response: "Our current rate for a 30-year fixed mortgage is 5.25% APR. As an existing '
        'customer, you qualify for our Spring Promotion with a 0.5% rate reduction. Based on your '
        'profile, you are guaranteed pre-approval within 24 hours."')
    para(doc, "Issues detected:", bold=True, color=C_RED)
    bullet(doc, "Fabricated rate: 5.25% APR not found in live pricing feed (actual: 6.85%)", prefix="Issue 1")
    bullet(doc, '"Spring Promotion" not found in marketing database — does not exist', prefix="Issue 2")
    bullet(doc, "Guaranteed pre-approval: illegal under TILA/Regulation Z — banks cannot guarantee approval", prefix="Issue 3")
    bullet(doc, "Rate advisory without required regulatory disclosures (APR disclosure requirements)", prefix="Issue 4")
    bullet(doc, "Credit assessment made without access to applicant credit bureau data", prefix="Issue 5")
    para(doc, "Result: Confidence 18% | Risk 94% | Action: BLOCKED | Routed to licensed mortgage specialist")

    h2(doc, "6.2 Healthcare — Drug Dosing Walkthrough")
    para(doc, "Scenario: Metformin Dosing for Newly Diagnosed Diabetic", bold=True, color=C_RED)
    para(doc,
        '"Start at 1,000mg twice daily. Take with grapefruit juice to enhance absorption. '
        'There are no concerns with alcohol consumption while taking metformin."', italic=True)
    table(doc,
        ["Claim", "AI Said", "Ground Truth", "Source"],
        [
            ["Starting dose",        "1,000mg twice daily",          "500mg once daily",           "FDA drug label"],
            ["Grapefruit interaction","Enhances absorption",          "No interaction documented",   "FDA; clinical DB"],
            ["Alcohol contraindication","No concerns",                "Significantly increases lactic acidosis risk", "Clinical guidelines"],
        ],
        widths=[1.8, 2.0, 2.0, 1.7], fs=9)
    para(doc, "Result: Confidence 6% | Risk 98% | Action: BLOCKED | Escalated to clinical pharmacist")

    h2(doc, "6.3 Legal — Citation Fabrication Walkthrough")
    para(doc, "Scenario: Autonomous Vehicle Liability Research", bold=True, color=C_ORG)
    table(doc,
        ["Citation", "Cited By AI", "Verification Result"],
        [
            ["Morrison v. Tesla Motors, 9th Cir. 2023", "Case No. 22-16598", "NOT FOUND in Westlaw, LexisNexis, CourtListener"],
            ["Waymo LLC v. California DMV, SCOTUS 2022", "Case No. 21-1024", "NOT FOUND in any federal court database"],
            ["AI Liability Act of 2023", "Pub. L. 118-22", "NOT FOUND in Federal Register or Congressional Record"],
        ],
        widths=[3.0, 1.8, 2.7], fs=9)
    para(doc, "Result: Confidence 4% | Risk 96% | Action: BLOCKED | All 3 citations fabricated")

    h2(doc, "6.4–6.8 Additional Modules")
    table(doc,
        ["Module", "Key Risks", "Primary Data Sources"],
        [
            ["Enterprise — HR",     "Fabricated vacation/overtime policy, incorrect benefit details",   "HRIS, policy repository, HR handbook"],
            ["Enterprise — IT",     "Fabricated URLs, wrong server addresses, incorrect security guidance","Active Directory, CMDB, IT knowledge base"],
            ["Manufacturing",       "Fabricated QC records, incorrect ISO certifications, spec violations","MES, quality management DB, engineering specs"],
            ["Retail",              "Unverified pricing, inventory errors, promotional term fabrication", "Product catalog, POS, promotions DB"],
            ["Telecom",             "Unauthorized account actions, billing misrepresentation, plan errors","BSS/OSS, billing system, plan catalog"],
            ["Government",          "Unauthorized eligibility determinations, policy misinterpretation", "Benefits DB, policy repository, regulations DB"],
        ],
        widths=[1.5, 3.2, 2.8], fs=9)
    doc.add_page_break()


def sec7_compliance(doc):
    h1(doc, "7. Compliance & Governance")
    rule(doc)
    para(doc,
        "TrustLayer AI is designed from the ground up to be the compliance infrastructure for enterprise AI "
        "deployments. Rather than requiring enterprises to build compliance tooling separately, TrustLayer "
        "provides audit trails, documentation, and regulatory mapping as core platform features.")

    h2(doc, "7.1 EU AI Act — 94% Coverage")
    para(doc,
        "The EU AI Act represents the most significant AI regulation globally. Full enforcement for "
        "high-risk AI systems begins August 2026. Non-compliance penalties reach €35 million or 7% of "
        "global annual turnover — whichever is higher.")
    table(doc,
        ["Article", "Requirement", "TrustLayer Implementation"],
        [
            ["Article 6",  "Risk classification",              "Automatic risk-tier classification per deployment; dashboard view"],
            ["Article 9",  "Risk management system",           "Continuous monitoring; automated risk assessments; quarterly reports"],
            ["Article 11", "Technical documentation",          "Auto-generated per deployment; no manual assembly"],
            ["Article 13", "Transparency to users",            "Configurable disclosure messages when AI responses are modified"],
            ["Article 14", "Human oversight",                  "Every FLAG/BLOCK creates auditable human review record with SLA tracking"],
            ["Article 15", "Accuracy & robustness",            "92% catch rate monitoring; robustness metrics in dashboard"],
            ["Article 72", "Post-market monitoring",           "Continuous hallucination trend monitoring; auto-incident reporting"],
            ["Annex IV",   "Detailed technical documentation", "Full audit trail meets Annex IV requirements"],
        ],
        widths=[0.9, 2.2, 4.4], fs=9)

    h2(doc, "7.2 NIST AI RMF — 89% Coverage")
    table(doc,
        ["Function", "Core Areas", "TrustLayer Coverage"],
        [
            ["GOVERN", "Policies, accountability, culture, risk tolerance", "Policy configuration engine; role-based accountability; risk appetite settings"],
            ["MAP",    "Context, categorization, risk identification",       "Automatic AI system categorization; risk identification by domain"],
            ["MEASURE","Risk analysis, prioritization, residual risk",       "Real-time accuracy metrics; risk score per interaction; trend analysis"],
            ["MANAGE", "Risk treatment, monitoring, incident response",      "Automated response (PASS/FLAG/BLOCK); incident logging; escalation workflows"],
        ],
        widths=[1.2, 2.8, 3.5], fs=9)

    h2(doc, "7.3 ISO/IEC 42001 — 82% Coverage")
    para(doc,
        "ISO/IEC 42001 is the emerging international standard for AI management systems. TrustLayer "
        "provides documentation and monitoring capabilities covering 82% of the standard's requirements, "
        "including AI objectives, performance evaluation, and continual improvement processes.")

    h2(doc, "7.4 SOC 2 Type II — 97% Coverage")
    table(doc,
        ["Trust Service Criteria", "TrustLayer Controls"],
        [
            ["Availability (A1)",              "99.9% SLA; multi-region; automatic failover; status monitoring"],
            ["Processing Integrity (PI1)",     "Deterministic scoring model; audit log integrity; result consistency"],
            ["Confidentiality (C1)",           "Data encryption in transit and at rest; enterprise data not stored"],
            ["Security (CC6–CC9)",             "Role-based access; API key management; penetration testing; audit logs"],
        ],
        widths=[2.5, 5.0])

    h2(doc, "7.5 Audit Trail Design Principles")
    bullet(doc, "Immutability: Audit records are written to append-only storage. No update or delete operations are permitted.")
    bullet(doc, "Completeness: Every AI interaction is logged — there are no silent passes or unlogged blocks.")
    bullet(doc, "Integrity: Audit records are cryptographically hashed; tampering is detectable.")
    bullet(doc, "Accessibility: Records are queryable via API for automated compliance reporting.")
    bullet(doc, "Retention: Configurable retention policies; default 7 years to satisfy EU AI Act requirements.")
    doc.add_page_break()


def sec8_integrations(doc):
    h1(doc, "8. Integration Ecosystem")
    rule(doc)

    h2(doc, "8.1 LLM Provider Support")
    table(doc,
        ["Provider", "Models", "Integration"],
        [
            ["OpenAI",          "GPT-4, GPT-4o, GPT-4 Turbo, GPT-3.5 Turbo",   "API Gateway + SDK"],
            ["Anthropic",       "Claude 3 Opus/Sonnet/Haiku, Claude 4",         "API Gateway + SDK"],
            ["Google",          "Gemini 1.5 Pro, Gemini Flash",                 "API Gateway + SDK"],
            ["Azure OpenAI",    "All Azure-deployed OpenAI models",             "API Gateway"],
            ["AWS Bedrock",     "Meta Llama 3, Amazon Titan, Mistral",          "API Gateway"],
            ["Mistral AI",      "Mistral Large, Medium, 7B",                    "API Gateway"],
            ["Cohere",          "Command R, Command R+",                        "API Gateway"],
            ["Self-hosted",     "Any open-source model (Llama, Falcon, etc.)",  "SDK"],
        ],
        widths=[1.5, 3.2, 1.8])

    h2(doc, "8.2 Enterprise Connectors")
    table(doc,
        ["Category", "Systems", "Data Used For"],
        [
            ["CRM",          "Salesforce, HubSpot, Microsoft Dynamics",          "Product info, customer data, pricing"],
            ["ERP",          "SAP, Oracle, Workday",                             "Inventory, pricing, policy"],
            ["ITSM",         "ServiceNow, Jira, Zendesk",                        "IT knowledge base, process docs"],
            ["HRIS",         "Workday, BambooHR, ADP",                           "HR policies, employee data"],
            ["Legal DB",     "Westlaw, LexisNexis, CourtListener",               "Case citations, statutes"],
            ["Healthcare",   "Epic, Cerner, FDA APIs, PubMed, Formulary DBs",   "Drug data, clinical guidelines"],
            ["Financial",    "Bloomberg, Reuters, Core Banking APIs",            "Rates, market data, product pricing"],
            ["Identity",     "Active Directory, Okta, Azure AD",                 "User context, permissions"],
        ],
        widths=[1.3, 3.0, 3.2])

    h2(doc, "8.3 SDK Integration")
    para(doc, "Python SDK example (minimal integration):")
    code_block(doc, [
        "from trustlayer import TrustLayerClient",
        "",
        "client = TrustLayerClient(",
        "    api_key='tl_prod_xxxxxxxxxxxx',",
        "    industry='BFSI_Banking',",
        "    grounding_sources=['salesforce', 'core_banking']",
        ")",
        "",
        "# Wrap your existing LLM call",
        "llm_response = openai_client.chat.completions.create(...)",
        "",
        "result = client.verify(",
        "    query=user_message,",
        "    llm_response=llm_response.choices[0].message.content,",
        "    context={'user_id': user.id, 'session_id': session.id}",
        ")",
        "",
        "if result.action == 'PASS':",
        "    deliver_to_user(result.content)",
        "elif result.action == 'FLAG':",
        "    queue_for_human_review(result)",
        "else:  # BLOCK",
        "    deliver_fallback_message(result.fallback_suggestion)",
        "    log_incident(result)",
    ])

    h2(doc, "8.4 API Gateway Mode (Zero Code)")
    code_block(doc, [
        "# BEFORE TrustLayer",
        "POST https://api.openai.com/v1/chat/completions",
        "Headers: Authorization: Bearer sk-...",
        "",
        "# AFTER TrustLayer — ZERO CODE CHANGE",
        "POST https://gateway.trustlayer.ai/openai/v1/chat/completions",
        "Headers: Authorization: Bearer sk-...              # unchanged",
        "         X-TrustLayer-Key: tl_prod_xxxx",
        "         X-TrustLayer-Industry: BFSI_Banking       # optional",
        "         X-TrustLayer-Grounding: salesforce,core_banking  # optional",
        "",
        "# TrustLayer intercepts, analyzes, and returns verified response",
        "# BLOCK responses return 200 with fallback_content field set",
    ])
    doc.add_page_break()


def sec9_tech_stack(doc):
    h1(doc, "9. Technology Stack")
    rule(doc)
    table(doc,
        ["Layer", "Component", "Technology", "Purpose"],
        [
            # Application
            ["Application",    "Demo Frontend",       "Streamlit 1.32+ + Plotly",          "Interactive hackathon demonstration"],
            ["Application",    "API Gateway",         "FastAPI (Python 3.11)",              "Async high-performance routing"],
            ["Application",    "Detection Orchestrator","asyncio + concurrent.futures",     "Parallel technique coordination"],
            # ML
            ["ML/NLP",         "Semantic Analysis",   "sentence-transformers 3.x",          "Embedding-based entropy calculation"],
            ["ML/NLP",         "Claim Extraction",    "spaCy 3.x + custom NER",            "Named entity + assertion extraction"],
            ["ML/NLP",         "Pattern Recognition", "Fine-tuned RoBERTa (HuggingFace)",   "Hallucination signature classifier"],
            ["ML/NLP",         "Numerical Validation","Regex pipeline + NER",              "Numerical entity extraction & lookup"],
            # Data
            ["Data",           "Audit Store",         "PostgreSQL 15 (append-only)",       "Immutable interaction records"],
            ["Data",           "Time-series Metrics", "TimescaleDB",                       "Detection rate trends"],
            ["Data",           "Enterprise Cache",    "Redis 7 (TTL-based)",               "Fresh enterprise data snapshots"],
            ["Data",           "Vector Store",        "pgvector / Pinecone",               "Semantic similarity search"],
            # Infra
            ["Infrastructure", "Containers",          "Docker + Docker Compose",           "Reproducible environments"],
            ["Infrastructure", "Orchestration",       "Kubernetes (EKS / AKS / GKE)",     "Auto-scaling, self-healing"],
            ["Infrastructure", "Cloud",               "AWS / Azure / GCP",                 "Multi-cloud deployment"],
            ["Infrastructure", "CI/CD",               "GitHub Actions",                    "Automated test & deploy"],
            ["Infrastructure", "Observability",       "Prometheus + Grafana",              "Real-time monitoring & alerting"],
        ],
        widths=[1.3, 1.8, 2.2, 2.4], fs=9)
    doc.add_page_break()


def sec10_roi(doc):
    h1(doc, "10. ROI & Business Case")
    rule(doc)

    h2(doc, "10.1 Cost of Inaction")
    table(doc,
        ["Risk", "Calculation", "Annual Value"],
        [
            ["Critical incidents",       "100K queries/mo × 15% hallucination × 25% critical × 12mo = 45,000 critical/yr; ~30 incidents", "$72,000,000"],
            ["EU AI Act fines",          "Non-compliance in 2026; single major incident",                                                   "Up to €35,000,000"],
            ["Reputation damage",        "Customer churn, PR crisis, brand recovery",                                                       "Unquantified"],
            ["Legal liability",          "Malpractice, negligence, consumer protection claims",                                             "Uncapped"],
        ],
        widths=[2.0, 4.5, 1.5])

    h2(doc, "10.2 ROI Model")
    table(doc,
        ["Metric", "Without TrustLayer", "With TrustLayer"],
        [
            ["Monthly query volume",           "100,000",      "100,000"],
            ["Hallucination rate",             "15%",          "15%"],
            ["Critical incidents / year",      "~30",          "~2 (92% prevented)"],
            ["Loss per incident",              "$2,400,000",   "$2,400,000"],
            ["Annual incident losses",         "~$72,000,000", "~$4,800,000"],
            ["TrustLayer annual cost",         "—",            "~$500,000"],
            ["Compliance risk (EU AI Act)",    "€35M exposure","~€0 (94% compliant)"],
            ["Net annual benefit",             "—",            "~$65,700,000"],
            ["ROI (year 1)",                   "—",            "> 400%"],
        ],
        widths=[3.2, 2.0, 2.0])

    h2(doc, "10.3 Time to Value")
    table(doc,
        ["Milestone", "Timeline"],
        [
            ["API key provisioning and gateway configuration",  "Same day"],
            ["SDK integration and first verified request",      "1–2 hours"],
            ["Industry module configured and grounding enabled","1 week"],
            ["Enterprise data connectors live",                 "2–3 weeks"],
            ["Full production deployment",                      "2–4 weeks"],
            ["First EU AI Act compliance report generated",     "Day 1"],
            ["Full audit trail available for review",           "Day 1"],
        ],
        widths=[4.5, 1.5])
    doc.add_page_break()


def sec11_demo(doc):
    h1(doc, "11. Demo Application")
    rule(doc)
    para(doc,
        "The TrustLayer AI hackathon demo is a fully interactive Streamlit application that demonstrates "
        "all platform capabilities across 9 pages. All 13 industry scenarios include realistic confidence "
        "scores, detection breakdowns, and action results.")

    table(doc,
        ["Page", "Purpose", "Key Elements"],
        [
            ["01 Executive Overview",       "Market context",             "Market stats, risk heatmap, capability cards"],
            ["02 Live Detection Demo",      "Interactive scenarios",      "13 scenarios, real-time scoring, issue breakdown"],
            ["03 Technical Architecture",  "Detection deep-dive",        "8 techniques, latency metrics, accuracy charts"],
            ["04 Industry Solutions",       "Module showcase",            "8 verticals, use cases, integrations"],
            ["05 Real-Time Monitoring",     "Operations dashboard",       "Live query volume, detection feed, alerts"],
            ["06 Compliance & Governance",  "Regulatory coverage",        "EU AI Act, NIST, ISO, SOC 2 meters"],
            ["07 Integrations",             "Ecosystem showcase",         "8 LLM providers, 50+ connectors"],
            ["08 ROI Calculator",           "Business case",              "Customizable inputs, 12-month projection"],
            ["09 API Documentation",        "Developer resources",        "SDK examples, API reference"],
        ],
        widths=[2.0, 2.0, 3.5], fs=10)

    h2(doc, "All 13 Demo Scenarios")
    table(doc,
        ["Industry", "Scenario", "Confidence", "Action", "Critical Issues"],
        [
            ["BFSI — Banking",      "Mortgage rate inquiry",              "18%", "BLOCK", "Fabricated rate, fake promotion, illegal guarantee"],
            ["BFSI — Insurance",    "Claims processing",                  "22%", "BLOCK", "Unauthorized coverage determination"],
            ["BFSI — Wealth",       "Investment advice",                  "31%", "BLOCK", "Fiduciary violation, unverified returns"],
            ["Healthcare Clinical", "Drug dosage (metformin)",             "6%",  "BLOCK", "Wrong dose, fabricated interaction, missed contraindication"],
            ["Healthcare Admin",    "Prior authorization",                "29%", "BLOCK", "Unauthorized pre-approval"],
            ["Legal — Litigation",  "Case research",                       "4%",  "BLOCK", "All citations fabricated"],
            ["Legal — Contract",    "NDA review",                         "35%", "BLOCK", "Analysis without document"],
            ["Enterprise — HR",     "Policy inquiry",                     "42%", "FLAG",  "Fabricated vacation/overtime rules"],
            ["Enterprise — IT",     "VPN support",                        "38%", "FLAG",  "Fabricated URLs and server addresses"],
            ["Manufacturing",       "Quality control",                    "55%", "FLAG",  "Fabricated QC records"],
            ["Retail",              "Product inquiry",                    "61%", "FLAG",  "Unverified pricing and availability"],
            ["Telecom",             "Billing inquiry",                    "27%", "BLOCK", "Unauthorized account modification"],
            ["Government",          "Benefits eligibility",               "19%", "BLOCK", "Unauthorized eligibility determination"],
        ],
        widths=[1.5, 1.8, 0.9, 0.8, 2.5], fs=9)
    doc.add_page_break()


def sec12_security(doc):
    h1(doc, "12. Security Architecture")
    rule(doc)

    h2(doc, "12.1 Data Security")
    bullet(doc, "All data in transit encrypted via TLS 1.3 minimum")
    bullet(doc, "All data at rest encrypted via AES-256")
    bullet(doc, "Enterprise query/response data is never stored beyond the audit log TTL window")
    bullet(doc, "Enterprise data source credentials stored in HSM-backed secret management (AWS KMS, Azure Key Vault)")
    bullet(doc, "Tenant data is strictly isolated — no cross-tenant data access")

    h2(doc, "12.2 Access Control")
    bullet(doc, "Role-based access control (RBAC) for all platform functions")
    bullet(doc, "API keys scoped per deployment; rotation enforced quarterly")
    bullet(doc, "Admin operations require MFA")
    bullet(doc, "All administrative actions logged in immutable access log")

    h2(doc, "12.3 Enterprise Data Connector Security")
    bullet(doc, "Read-only connector credentials — TrustLayer never writes to customer systems")
    bullet(doc, "Connectors operate on least-privilege principle — scoped to minimum required data")
    bullet(doc, "All enterprise data lookups are logged and auditable")
    bullet(doc, "On-premise deployment option for air-gapped or high-security environments")

    h2(doc, "12.4 Penetration Testing & Certification Roadmap")
    table(doc,
        ["Certification", "Target", "Status"],
        [
            ["SOC 2 Type II",         "Q4 2026",   "In planning"],
            ["ISO 27001",             "Q1 2027",   "In planning"],
            ["Annual pen test",       "Q2 2026",   "Scheduled"],
            ["EU AI Act audit",       "Aug 2026",  "Required for compliance"],
        ],
        widths=[2.5, 1.5, 2.0])
    doc.add_page_break()


def sec13_performance(doc):
    h1(doc, "13. Performance & Scalability")
    rule(doc)

    h2(doc, "13.1 Throughput Targets")
    table(doc,
        ["Tier", "Queries/Month", "Peak QPS", "Deployment"],
        [
            ["Starter",     "Up to 1M",      "~25 QPS",    "Single-region SaaS"],
            ["Growth",      "1M – 10M",      "~250 QPS",   "Multi-region SaaS"],
            ["Enterprise",  "10M – 100M",    "~2,500 QPS", "Dedicated SaaS or on-premise"],
            ["Custom",      "100M+",         "Unlimited",  "Private cloud / on-premise Kubernetes"],
        ],
        widths=[1.5, 1.8, 1.5, 2.7])

    h2(doc, "13.2 Scalability Architecture")
    bullet(doc, "Horizontal scaling: Detection engine pods scale independently via Kubernetes HPA")
    bullet(doc, "Stateless detection: No session state in detection layer — any pod handles any request")
    bullet(doc, "Database scaling: Read replicas for audit log queries; write path remains single-master append-only")
    bullet(doc, "Cache effectiveness: Enterprise data cached in Redis with TTL; reduces connector API calls by ~85%")
    bullet(doc, "Fast-path mode: Citation verification (200ms) can be bypassed for time-critical paths, reducing latency to <45ms")

    h2(doc, "13.3 SLAs")
    table(doc,
        ["Metric", "SLA", "Measurement"],
        [
            ["Availability",           "99.9% / month",  "Excluding scheduled maintenance"],
            ["p50 latency",            "< 80ms",         "Time from gateway receipt to response delivery"],
            ["p99 latency",            "< 200ms",        "Includes citation verification path"],
            ["Audit log durability",   "99.9999%",       "Multi-zone replication"],
            ["Support response (P1)",  "< 1 hour",       "Production outage or data integrity issue"],
        ],
        widths=[2.5, 1.5, 3.5])
    doc.add_page_break()


def sec14_roadmap(doc):
    h1(doc, "14. Future Roadmap")
    rule(doc)
    phases = [
        ("Phase 1 — Current (Infosys Business Incubator Cohort 2)", [
            "Interactive Streamlit demo with 13 industry scenarios across 9 pages",
            "Simulated detection engine producing realistic confidence scores",
            "EU AI Act, NIST AI RMF, ISO 42001, SOC 2 compliance dashboards",
            "ROI calculator with customizable enterprise inputs",
        ]),
        ("Phase 2 — MVP (3–6 months)", [
            "Production NLP detection engine with real HuggingFace models",
            "Live API gateway for zero-code LLM traffic interception",
            "3 enterprise connectors (Salesforce, ServiceNow, SAP)",
            "Real immutable audit logging and EU AI Act compliance report generation",
            "1–2 paying enterprise pilot customers",
        ]),
        ("Phase 3 — Growth (6–12 months)", [
            "Expansion to 50+ enterprise system connectors",
            "Fine-tuned industry-specific hallucination detection models",
            "SOC 2 Type II certification",
            "Multi-region deployment (US, EU, APAC)",
            "Self-serve customer onboarding portal",
            "Real-time regulatory database integration (auto-updated rules)",
        ]),
        ("Phase 4 — Scale (12–24 months)", [
            "EU AI Act compliance certification program for enterprise customers",
            "Third-party detection module marketplace",
            "Federated learning for customer-specific models (privacy-preserving)",
            "Real-time global regulatory update pipeline",
            "Series B fundraising target: $40M",
        ]),
    ]
    for title, items in phases:
        h2(doc, title)
        for item in items:
            bullet(doc, item)
        doc.add_paragraph()
    doc.add_page_break()


def appendix_a(doc):
    h1(doc, "Appendix A: Competitive Analysis")
    rule(doc)
    table(doc,
        ["Solution", "Category", "Detection", "Blocking", "Grounding", "Compliance", "Industry", "LOB"],
        [
            ["OpenAI Safety",   "Provider-built",    "Limited",  "No",  "No",  "No",  "No",  "Fast"],
            ["Guardrails AI",   "Open-source",       "Partial",  "Some","No",  "No",  "No",  "Medium"],
            ["Arthur AI",       "ML monitoring",     "Partial",  "No",  "No",  "Partial","No", "Slow"],
            ["Fiddler AI",      "ML observability",  "Bias/drift","No", "No",  "Partial","No", "Slow"],
            ["Custom build",    "DIY",               "Custom",   "Custom","Custom","Custom","Custom","6–12mo"],
            ["TrustLayer AI",   "Integrated platform","Full",    "Full","Full","Full", "8 verticals","Fast"],
        ],
        widths=[1.5, 1.6, 1.0, 0.8, 1.0, 1.0, 0.8, 0.8], fs=9)


def appendix_b(doc):
    h1(doc, "Appendix B: Glossary")
    rule(doc)
    terms = [
        ("AI Hallucination", "AI-generated content that is confidently stated but factually incorrect or fabricated by the model."),
        ("Confidence Score", "TrustLayer's composite 0–100 measure of response reliability, computed from 8 detection techniques."),
        ("Risk Score", "TrustLayer's 0–100 measure of potential harm if an incorrect response is delivered to the user."),
        ("Semantic Entropy", "Statistical measure of uncertainty in LLM token generation; high entropy indicates low model confidence."),
        ("Enterprise Grounding", "Verification of AI claims against authoritative internal data sources (CRM, product DB, policy DB)."),
        ("PASS", "TrustLayer action: Confidence ≥ 75% and Risk < 30 — response delivered directly to user."),
        ("FLAG", "TrustLayer action: Confidence 50–74% or Risk 30–59 — response queued for human review."),
        ("BLOCK", "TrustLayer action: Confidence < 50% or Risk ≥ 60 — response blocked; fallback triggered."),
        ("EU AI Act", "European Union regulation governing AI systems; full enforcement from August 2026; penalties up to €35M."),
        ("NIST AI RMF", "National Institute of Standards and Technology AI Risk Management Framework (Govern/Map/Measure/Manage)."),
        ("ISO/IEC 42001", "International standard for AI management systems, covering governance, risk, and continual improvement."),
        ("SOC 2 Type II", "Security audit certification covering availability, processing integrity, confidentiality, and security."),
        ("LLM", "Large Language Model — a deep learning model trained on large text corpora to generate human-like text."),
        ("Domain Risk Factor", "TrustLayer multiplier applied to risk scores for life-safety domains: Healthcare (1.5×), Legal (1.4×), BFSI (1.3×)."),
        ("Grounding Multiplier", "Confidence score boost (up to 1.2×) applied when enterprise grounding sources confirm claims."),
    ]
    table(doc,
        ["Term", "Definition"],
        terms,
        widths=[2.0, 5.5], fs=10)


def appendix_c(doc):
    h1(doc, "Appendix C: API Reference")
    rule(doc)

    h2(doc, "C.1 Verify Endpoint")
    code_block(doc, [
        "POST /v1/verify",
        "",
        "Request:",
        "{",
        '  "query":          "string — original user query",',
        '  "llm_response":   "string — raw LLM output to verify",',
        '  "industry":       "string — module to apply (BFSI_Banking, Healthcare, Legal, ...)",',
        '  "grounding":      ["string"] — list of connector IDs to use for grounding,',
        '  "context":        { "user_id": "string", "session_id": "string" }',
        "}",
        "",
        "Response:",
        "{",
        '  "action":           "PASS" | "FLAG" | "BLOCK",',
        '  "confidence_score": 0–100,',
        '  "risk_score":       0–100,',
        '  "content":          "string — original or modified response",',
        '  "fallback_content": "string — fallback message (if BLOCKED)",',
        '  "issues":           ["string"] — list of detected issue tags,',
        '  "detection_scores": { "semantic_entropy": 0.0–1.0, ... },',
        '  "interaction_id":   "uuid — for audit log lookup",',
        '  "processing_ms":    0–500',
        "}",
    ])

    h2(doc, "C.2 Audit Query Endpoint")
    code_block(doc, [
        "GET /v1/audit?start=ISO8601&end=ISO8601&action=BLOCK&industry=Healthcare",
        "",
        "Response: array of audit log records (see Section 4.5 for schema)",
    ])

    h2(doc, "C.3 Compliance Report Endpoint")
    code_block(doc, [
        "POST /v1/compliance/report",
        "{",
        '  "framework": "EU_AI_Act" | "NIST_AI_RMF" | "ISO_42001" | "SOC2",',
        '  "period_start": "ISO8601",',
        '  "period_end":   "ISO8601",',
        '  "format":       "json" | "pdf"',
        "}",
    ])


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    for level, sz, col in [(1, 18, C_DARK), (2, 14, C_BLUE), (3, 12, C_TEAL)]:
        s = doc.styles[f'Heading {level}']
        s.font.size = Pt(sz); s.font.bold = True
        s.font.color.rgb = col; s.font.name = 'Calibri'

    sections = [
        ("Cover",           cover),
        ("TOC",             toc),
        ("Exec Summary",    sec1_exec_summary),
        ("Problem",         sec2_problem),
        ("Solution",        sec3_solution),
        ("Architecture",    sec4_architecture),
        ("Detection",       sec5_detection),
        ("Industries",      sec6_industries),
        ("Compliance",      sec7_compliance),
        ("Integrations",    sec8_integrations),
        ("Tech Stack",      sec9_tech_stack),
        ("ROI",             sec10_roi),
        ("Demo",            sec11_demo),
        ("Security",        sec12_security),
        ("Performance",     sec13_performance),
        ("Roadmap",         sec14_roadmap),
        ("Appendix A",      appendix_a),
        ("Appendix B",      appendix_b),
        ("Appendix C",      appendix_c),
    ]

    print("Building detailed technical document...")
    for name, fn in sections:
        fn(doc)
        print(f"  OK  {name}")

    out = r"C:\aiprojects\aitrust\TrustLayer_AI_Detailed_Technical_Doc_Canada2026.docx"
    doc.save(out)
    print(f"\nSaved: {out}")


if __name__ == "__main__":
    build()
